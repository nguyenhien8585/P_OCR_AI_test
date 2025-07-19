import streamlit as st
import requests
import base64
import io
import json
from PIL import Image, ImageDraw, ImageFilter, ImageEnhance
import fitz  # PyMuPDF
import tempfile
import os
import re
import time
import math
import gc  # Garbage collection

# Import python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import cv2
    import numpy as np
    from scipy import ndimage
    from skimage import filters, measure, morphology
    CV2_AVAILABLE = True
except ImportError:
    CV2_AVAILABLE = False

# Cấu hình trang
st.set_page_config(
    page_title="PDF/LaTeX Converter - Enhanced with Mistral OCR & Phone Processing",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# CSS cải tiến
try:
    st.markdown("""
    <style>
        .main-header {
            text-align: center;
            color: #2E86AB;
            font-size: 2.5rem;
            margin-bottom: 2rem;
            text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
        }
        
        .latex-output {
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 1.5rem;
            border-radius: 12px;
            font-family: 'Consolas', 'Monaco', monospace;
            border-left: 4px solid #2E86AB;
            max-height: 400px;
            overflow-y: auto;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        }
        
        .extracted-image {
            border: 3px solid #28a745;
            border-radius: 12px;
            margin: 15px 0;
            padding: 10px;
            background: linear-gradient(135deg, #f8f9fa 0%, #ffffff 100%);
            box-shadow: 0 6px 12px rgba(0,0,0,0.15);
            transition: transform 0.3s ease;
        }
        
        .extracted-image:hover {
            transform: translateY(-5px);
            box-shadow: 0 8px 16px rgba(0,0,0,0.2);
        }
        
        .metric-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 1.5rem;
            border-radius: 12px;
            text-align: center;
            margin: 8px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.15);
            transition: transform 0.2s ease;
        }
        
        .metric-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 6px 12px rgba(0,0,0,0.2);
        }
        
        .figure-preview {
            border: 2px solid #007bff;
            border-radius: 8px;
            padding: 8px;
            margin: 8px 0;
            background: white;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }
        
        .figure-info {
            background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
            padding: 0.8rem;
            border-radius: 6px;
            margin: 5px 0;
            font-size: 0.85rem;
            border-left: 3px solid #ffc107;
        }
        
        .status-success {
            background: linear-gradient(135deg, #d4edda 0%, #c3e6cb 100%);
            color: #155724;
            padding: 1rem;
            border-radius: 8px;
            border-left: 4px solid #28a745;
            margin: 10px 0;
        }
        
        .status-warning {
            background: linear-gradient(135deg, #fff3cd 0%, #ffeaa7 100%);
            color: #856404;
            padding: 1rem;
            border-radius: 8px;
            border-left: 4px solid #ffc107;
            margin: 10px 0;
        }
        
        .processing-container {
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 2rem;
            border-radius: 12px;
            margin: 20px 0;
            border: 2px solid #dee2e6;
        }

        .mistral-badge {
            background: linear-gradient(135deg, #FF6B35 0%, #FF8E53 100%);
            color: white;
            padding: 0.5rem 1rem;
            border-radius: 20px;
            font-weight: bold;
            text-align: center;
            margin: 5px 0;
        }
    </style>
    """, unsafe_allow_html=True)
except Exception as e:
    st.error(f"CSS loading error: {str(e)}")

class MistralOCRService:
    """
    Mistral OCR Service để đếm figures trong ảnh và phân tích nội dung
    """
    
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://api.mistral.ai/v1/chat/completions"
        self.session = requests.Session()
        self.session.headers.update({
            'Content-Type': 'application/json',
            'Authorization': f'Bearer {api_key}',
            'User-Agent': 'PDF-LaTeX-Converter/1.0'
        })
        self.max_retries = 3
        self.timeout = 60
    
    def analyze_image_content(self, image_bytes, detect_figures=True, detect_tables=True):
        """
        Phân tích nội dung ảnh và đếm số lượng figures/tables bằng Mistral Vision
        """
        try:
            # Encode image
            encoded_image = base64.b64encode(image_bytes).decode('utf-8')
            
            # Tạo prompt để phân tích figures
            analysis_prompt = f"""
Analyze this image carefully and count the number of figures, diagrams, charts, tables, and illustrations.

Please provide a JSON response with the following structure:
{{
    "figure_count": <number_of_figures_diagrams_charts>,
    "table_count": <number_of_tables>,
    "total_count": <total_visual_elements>,
    "confidence": <confidence_score_0_to_1>,
    "analysis": {{
        "has_mathematical_content": <true/false>,
        "has_text_content": <true/false>,
        "content_type": "<exam/textbook/document/mixed>",
        "visual_complexity": "<low/medium/high>",
        "layout_type": "<single_column/multi_column/mixed>"
    }},
    "visual_elements": [
        {{
            "type": "<figure/table/diagram/chart>",
            "description": "<brief_description>",
            "estimated_position": "<top/middle/bottom>"
        }}
    ]
}}

Instructions:
- Count ALL visual elements that are NOT just plain text
- Include: graphs, charts, diagrams, illustrations, tables, mathematical figures
- Exclude: plain text paragraphs, headers, footers
- Be conservative but accurate in counting
- Provide confidence score based on image clarity and your certainty
"""

            # Chuẩn bị payload cho Mistral
            payload = {
                "model": "pixtral-12b-2409",  # Mistral's vision model
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "text",
                                "text": analysis_prompt
                            },
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{encoded_image}"
                                }
                            }
                        ]
                    }
                ],
                "temperature": 0.1,
                "max_tokens": 1000,
                "response_format": {
                    "type": "json_object"
                }
            }
            
            # Call API với retry logic
            for attempt in range(self.max_retries):
                try:
                    response = self.session.post(
                        self.base_url,
                        json=payload,
                        timeout=self.timeout
                    )
                    
                    if response.status_code == 200:
                        result = response.json()
                        if 'choices' in result and len(result['choices']) > 0:
                            content = result['choices'][0]['message']['content']
                            return self._process_mistral_response(content)
                        else:
                            if attempt == self.max_retries - 1:
                                st.warning("⚠️ Mistral API không trả về kết quả hợp lệ - sử dụng fallback")
                                return self._get_fallback_result()
                    elif response.status_code == 401:
                        st.error("❌ Mistral API key không hợp lệ")
                        return self._get_fallback_result()
                    elif response.status_code == 429:
                        if attempt < self.max_retries - 1:
                            time.sleep(2 ** attempt)
                            continue
                        st.warning("⚠️ Mistral API rate limit - sử dụng fallback")
                        return self._get_fallback_result()
                    else:
                        if attempt == self.max_retries - 1:
                            st.warning(f"⚠️ Mistral API error {response.status_code} - sử dụng fallback")
                            return self._get_fallback_result()
                        
                except requests.exceptions.Timeout:
                    if attempt == self.max_retries - 1:
                        st.warning("⚠️ Mistral API timeout - sử dụng fallback")
                        return self._get_fallback_result()
                    time.sleep(2 ** attempt)
                    continue
                except Exception as e:
                    if attempt == self.max_retries - 1:
                        st.warning(f"⚠️ Mistral API error: {str(e)} - sử dụng fallback")
                        return self._get_fallback_result()
                    time.sleep(2 ** attempt)
                    continue
            
            return self._get_fallback_result()
            
        except Exception as e:
            st.warning(f"⚠️ Mistral OCR error: {str(e)} - sử dụng fallback method")
            return self._get_fallback_result()
    
    def count_figures_in_text(self, text_content):
        """
        Đếm số lượng figures được nhắc đến trong text bằng Mistral
        """
        try:
            payload = {
                "model": "mistral-small-latest",
                "messages": [
                    {
                        "role": "user",
                        "content": f"""
Analyze this text and count how many figures, tables, diagrams, or visual elements are mentioned or referenced.

Text: {text_content[:2000]}

Please respond with a JSON object:
{{
    "figure_count": <number>,
    "table_count": <number>,
    "references": ["list of figure/table references found"]
}}

Look for patterns like: "hình", "figure", "fig", "bảng", "table", "biểu đồ", "đồ thị", "chart", "diagram", etc.
"""
                    }
                ],
                "temperature": 0.1,
                "max_tokens": 200,
                "response_format": {
                    "type": "json_object"
                }
            }
            
            response = self.session.post(self.base_url, json=payload, timeout=30)
            
            if response.status_code == 200:
                result = response.json()
                if 'choices' in result and len(result['choices']) > 0:
                    content = json.loads(result['choices'][0]['message']['content'])
                    return content.get('figure_count', 0), content.get('table_count', 0)
            
            return 0, 0
            
        except Exception:
            return 0, 0
    
    def _process_mistral_response(self, response_content):
        """
        Xử lý response từ Mistral API
        """
        try:
            data = json.loads(response_content)
            
            # Extract counts với fallback values
            figure_count = data.get('figure_count', 0)
            table_count = data.get('table_count', 0)
            total_count = data.get('total_count', figure_count + table_count)
            confidence = data.get('confidence', 0.8)
            
            # Extract analysis info
            analysis = data.get('analysis', {})
            visual_elements = data.get('visual_elements', [])
            
            # Convert confidence to 0-1 range if needed
            if confidence > 1:
                confidence = confidence / 100
            
            return {
                'success': True,
                'figure_count': max(0, int(figure_count)),
                'table_count': max(0, int(table_count)),
                'total_count': max(1, int(total_count)),  # At least 1
                'confidence': min(1.0, max(0.0, float(confidence))),
                'analysis': analysis,
                'visual_elements': visual_elements,
                'method': 'mistral_ocr',
                'model': 'pixtral-12b-2409'
            }
            
        except json.JSONDecodeError:
            # Try to extract numbers from text response
            try:
                import re
                figure_matches = re.findall(r'figure[s]?["\s]*:?\s*(\d+)', response_content, re.IGNORECASE)
                table_matches = re.findall(r'table[s]?["\s]*:?\s*(\d+)', response_content, re.IGNORECASE)
                
                figure_count = int(figure_matches[0]) if figure_matches else 2
                table_count = int(table_matches[0]) if table_matches else 1
                
                return {
                    'success': True,
                    'figure_count': figure_count,
                    'table_count': table_count,
                    'total_count': figure_count + table_count,
                    'confidence': 0.7,
                    'method': 'mistral_ocr_fallback',
                    'analysis': {},
                    'visual_elements': []
                }
            except:
                return self._get_fallback_result()
        except Exception:
            return self._get_fallback_result()
    
    def _get_fallback_result(self):
        """
        Fallback result khi Mistral API không khả dụng
        """
        return {
            'success': False,
            'figure_count': 2,  # Conservative estimate
            'table_count': 1,
            'total_count': 3,
            'confidence': 0.5,
            'method': 'fallback',
            'analysis': {},
            'visual_elements': []
        }

class PhoneImageProcessor:
    """
    Xử lý ảnh chụp từ điện thoại để tối ưu cho OCR - Enhanced Version
    """
    
    @staticmethod
    def process_phone_image(image_bytes, auto_enhance=True, auto_rotate=True, 
                          perspective_correct=True, text_enhance=True, 
                          crop_document=True, noise_reduction=True):
        """
        Xử lý ảnh điện thoại với các tùy chọn nâng cao
        """
        try:
            # Đọc ảnh
            img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            
            # Convert to numpy for CV2 processing if available
            if CV2_AVAILABLE:
                img = np.array(img_pil)
                original_img = img.copy()
                
                # Step 1: Noise reduction (if enabled)
                if noise_reduction:
                    img = PhoneImageProcessor._reduce_noise(img)
                
                # Step 2: Document detection and cropping
                if crop_document:
                    img = PhoneImageProcessor._smart_document_crop(img)
                
                # Step 3: Auto rotate & straighten
                if auto_rotate:
                    img = PhoneImageProcessor._enhanced_auto_rotate(img)
                
                # Step 4: Perspective correction
                if perspective_correct:
                    img = PhoneImageProcessor._enhanced_perspective_correction(img)
                
                # Step 5: Auto enhance
                if auto_enhance:
                    img = PhoneImageProcessor._enhanced_auto_enhance(img)
                
                # Step 6: Text enhancement
                if text_enhance:
                    img = PhoneImageProcessor._enhanced_text_enhancement(img)
                
                # Convert back to PIL
                processed_img = Image.fromarray(img)
            else:
                # Fallback: basic PIL processing
                processed_img = img_pil
                
                if auto_enhance:
                    # Basic enhancement with PIL
                    from PIL import ImageEnhance
                    enhancer = ImageEnhance.Contrast(processed_img)
                    processed_img = enhancer.enhance(1.3)
                    
                    enhancer = ImageEnhance.Sharpness(processed_img)
                    processed_img = enhancer.enhance(1.2)
                    
                    enhancer = ImageEnhance.Brightness(processed_img)
                    processed_img = enhancer.enhance(1.1)
            
            return processed_img
            
        except Exception as e:
            st.error(f"❌ Lỗi xử lý ảnh: {str(e)}")
            return Image.open(io.BytesIO(image_bytes)).convert("RGB")
    
    @staticmethod
    def _reduce_noise(img):
        """Giảm noise trong ảnh"""
        try:
            # Bilateral filter để giảm noise mà vẫn giữ edges
            denoised = cv2.bilateralFilter(img, 9, 75, 75)
            return denoised
        except Exception:
            return img
    
    @staticmethod
    def _smart_document_crop(img):
        """Tự động crop document thông minh"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            h, w = gray.shape
            
            # Enhanced edge detection
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            edges = cv2.Canny(blurred, 30, 80, apertureSize=3)
            
            # Morphological operations để connect broken lines
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
            edges = cv2.morphologyEx(edges, cv2.MORPH_CLOSE, kernel)
            
            # Find contours
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            contours = sorted(contours, key=cv2.contourArea, reverse=True)
            
            # Look for document-like contours
            for contour in contours[:10]:
                # Approximate contour
                epsilon = 0.02 * cv2.arcLength(contour, True)
                approx = cv2.approxPolyDP(contour, epsilon, True)
                
                # Check if it's roughly rectangular (4-8 points)
                if 4 <= len(approx) <= 8:
                    area = cv2.contourArea(contour)
                    img_area = h * w
                    area_ratio = area / img_area
                    
                    # Must be substantial portion of image
                    if 0.1 <= area_ratio <= 0.95:
                        # Get bounding rectangle
                        x, y, w_rect, h_rect = cv2.boundingRect(contour)
                        
                        # Add some padding
                        padding = 20
                        x = max(0, x - padding)
                        y = max(0, y - padding)
                        w_rect = min(w - x, w_rect + 2*padding)
                        h_rect = min(h - y, h_rect + 2*padding)
                        
                        # Crop the image
                        cropped = img[y:y+h_rect, x:x+w_rect]
                        
                        # Validate crop
                        if cropped.shape[0] > 100 and cropped.shape[1] > 100:
                            return cropped
            
            return img
            
        except Exception:
            return img
    
    @staticmethod
    def _enhanced_auto_rotate(img):
        """Tự động xoay ảnh thông minh hơn"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            
            # Method 1: Hough lines
            edges = cv2.Canny(gray, 50, 150, apertureSize=3)
            lines = cv2.HoughLines(edges, 1, np.pi/180, threshold=80)
            
            angles = []
            if lines is not None:
                for rho, theta in lines[:20]:  # More lines for better accuracy
                    angle = theta * 180 / np.pi
                    # Normalize angle to [-45, 45]
                    if angle > 90:
                        angle = angle - 180
                    elif angle > 45:
                        angle = angle - 90
                    elif angle < -45:
                        angle = angle + 90
                    
                    if abs(angle) < 45:  # Filter extreme angles
                        angles.append(angle)
            
            if angles:
                # Use median for robustness
                rotation_angle = np.median(angles)
                
                # Only rotate if angle is significant
                if abs(rotation_angle) > 0.5:
                    center = (img.shape[1]//2, img.shape[0]//2)
                    M = cv2.getRotationMatrix2D(center, rotation_angle, 1.0)
                    
                    # Calculate new image size to avoid cropping
                    cos = np.abs(M[0, 0])
                    sin = np.abs(M[0, 1])
                    new_w = int((img.shape[0] * sin) + (img.shape[1] * cos))
                    new_h = int((img.shape[0] * cos) + (img.shape[1] * sin))
                    
                    # Adjust transformation matrix
                    M[0, 2] += (new_w / 2) - center[0]
                    M[1, 2] += (new_h / 2) - center[1]
                    
                    img = cv2.warpAffine(img, M, (new_w, new_h), 
                                       flags=cv2.INTER_CUBIC, 
                                       borderMode=cv2.BORDER_CONSTANT,
                                       borderValue=(255, 255, 255))
            
            return img
            
        except Exception:
            return img
    
    @staticmethod
    def _enhanced_perspective_correction(img):
        """Sửa perspective distortion nâng cao"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            h, w = gray.shape
            
            # Multiple methods for document detection
            
            # Method 1: Adaptive thresholding + morphology
            adaptive_thresh = cv2.adaptiveThreshold(gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                                  cv2.THRESH_BINARY, 11, 2)
            
            # Method 2: Enhanced edge detection
            blurred = cv2.GaussianBlur(gray, (5, 5), 0)
            edges = cv2.Canny(blurred, 50, 150, apertureSize=3)
            
            # Combine both methods
            combined = cv2.bitwise_or(edges, cv2.bitwise_not(adaptive_thresh))
            
            # Morphological operations
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
            combined = cv2.morphologyEx(combined, cv2.MORPH_CLOSE, kernel)
            
            # Find contours
            contours, _ = cv2.findContours(combined, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            contours = sorted(contours, key=cv2.contourArea, reverse=True)
            
            # Look for document contour
            for contour in contours[:5]:
                peri = cv2.arcLength(contour, True)
                approx = cv2.approxPolyDP(contour, 0.015 * peri, True)  # More flexible approximation
                
                area = cv2.contourArea(contour)
                img_area = h * w
                area_ratio = area / img_area
                
                # Check for document-like properties
                if (len(approx) >= 4 and area_ratio > 0.2):
                    # If more than 4 points, find the best 4 corners
                    if len(approx) > 4:
                        # Use convex hull and find extreme points
                        hull = cv2.convexHull(contour)
                        
                        # Find the 4 extreme points
                        pts = hull.reshape(-1, 2)
                        
                        # Find corners
                        def distance(p1, p2):
                            return np.sqrt((p1[0]-p2[0])**2 + (p1[1]-p2[1])**2)
                        
                        # Find 4 corners by finding points that are farthest from each other
                        corners = []
                        
                        # Top-left: minimum sum
                        tl = pts[np.argmin(pts.sum(axis=1))]
                        corners.append(tl)
                        
                        # Bottom-right: maximum sum  
                        br = pts[np.argmax(pts.sum(axis=1))]
                        corners.append(br)
                        
                        # Top-right: minimum diff (x-y)
                        tr = pts[np.argmin(np.diff(pts, axis=1).flatten())]
                        corners.append(tr)
                        
                        # Bottom-left: maximum diff (x-y)
                        bl = pts[np.argmax(np.diff(pts, axis=1).flatten())]
                        corners.append(bl)
                        
                        approx = np.array(corners)
                    
                    if len(approx) == 4:
                        # Order points properly
                        rect = PhoneImageProcessor._order_points_enhanced(approx.reshape(-1, 2))
                        
                        # Calculate perspective transform
                        (tl, tr, br, bl) = rect
                        
                        # Calculate the width and height of the new image
                        widthA = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
                        widthB = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
                        maxWidth = max(int(widthA), int(widthB))
                        
                        heightA = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
                        heightB = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
                        maxHeight = max(int(heightA), int(heightB))
                        
                        # Ensure reasonable dimensions
                        if maxWidth > 100 and maxHeight > 100:
                            # Destination points
                            dst = np.array([
                                [0, 0],
                                [maxWidth - 1, 0],
                                [maxWidth - 1, maxHeight - 1],
                                [0, maxHeight - 1]], dtype="float32")
                            
                            # Apply perspective transformation
                            M = cv2.getPerspectiveTransform(rect, dst)
                            warped = cv2.warpPerspective(img, M, (maxWidth, maxHeight))
                            
                            return warped
            
            return img
            
        except Exception:
            return img
    
    @staticmethod
    def _order_points_enhanced(pts):
        """Enhanced point ordering"""
        # Sort points based on their x+y values (top-left has smallest sum)
        rect = np.zeros((4, 2), dtype="float32")
        
        # Top-left point has the smallest sum
        # Bottom-right point has the largest sum
        s = pts.sum(axis=1)
        rect[0] = pts[np.argmin(s)]
        rect[2] = pts[np.argmax(s)]
        
        # Top-right point has the smallest difference
        # Bottom-left point has the largest difference
        diff = np.diff(pts, axis=1)
        rect[1] = pts[np.argmin(diff)]
        rect[3] = pts[np.argmax(diff)]
        
        return rect
    
    @staticmethod
    def _enhanced_auto_enhance(img):
        """Tự động tăng cường chất lượng ảnh nâng cao"""
        try:
            # Method 1: CLAHE on LAB color space
            lab = cv2.cvtColor(img, cv2.COLOR_RGB2LAB)
            l, a, b = cv2.split(lab)
            
            # Apply CLAHE to L channel with optimized parameters
            clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
            l = clahe.apply(l)
            
            # Merge back
            enhanced = cv2.merge([l, a, b])
            enhanced = cv2.cvtColor(enhanced, cv2.COLOR_LAB2RGB)
            
            # Method 2: Gamma correction for brightness
            gamma = PhoneImageProcessor._calculate_optimal_gamma(enhanced)
            enhanced = PhoneImageProcessor._apply_gamma_correction(enhanced, gamma)
            
            # Method 3: Contrast enhancement
            enhanced = PhoneImageProcessor._enhance_contrast_adaptive(enhanced)
            
            return enhanced
            
        except Exception:
            return img
    
    @staticmethod
    def _calculate_optimal_gamma(img):
        """Tính gamma tối ưu dựa trên histogram"""
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            mean_brightness = np.mean(gray)
            
            # Gamma correction based on image brightness
            if mean_brightness < 100:  # Dark image
                return 0.7
            elif mean_brightness > 180:  # Bright image
                return 1.3
            else:  # Normal image
                return 1.0
        except:
            return 1.0
    
    @staticmethod
    def _apply_gamma_correction(img, gamma):
        """Áp dụng gamma correction"""
        try:
            # Build lookup table
            inv_gamma = 1.0 / gamma
            table = np.array([((i / 255.0) ** inv_gamma) * 255 for i in np.arange(0, 256)]).astype("uint8")
            
            # Apply gamma correction
            return cv2.LUT(img, table)
        except:
            return img
    
    @staticmethod
    def _enhance_contrast_adaptive(img):
        """Tăng cường contrast adaptive"""
        try:
            # Convert to YUV color space
            yuv = cv2.cvtColor(img, cv2.COLOR_RGB2YUV)
            
            # Apply histogram equalization to Y channel
            yuv[:,:,0] = cv2.equalizeHist(yuv[:,:,0])
            
            # Convert back to RGB
            enhanced = cv2.cvtColor(yuv, cv2.COLOR_YUV2RGB)
            
            return enhanced
        except:
            return img
    
    @staticmethod
    def _enhanced_text_enhancement(img):
        """Tăng cường text nâng cao"""
        try:
            # Convert to grayscale for processing
            if len(img.shape) == 3:
                gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            else:
                gray = img.copy()
            
            # Method 1: Advanced unsharp masking
            gaussian_3 = cv2.GaussianBlur(gray, (0, 0), 2.0)
            unsharp_mask = cv2.addWeighted(gray, 2.0, gaussian_3, -1.0, 0)
            
            # Method 2: High-pass filter
            kernel = np.array([[-1,-1,-1], [-1,9,-1], [-1,-1,-1]])
            sharpened = cv2.filter2D(unsharp_mask, -1, kernel)
            
            # Method 3: Morphological operations for text cleanup
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 1))
            cleaned = cv2.morphologyEx(sharpened, cv2.MORPH_CLOSE, kernel)
            
            # Method 4: Adaptive thresholding for binarization (optional)
            # This can help with very poor quality text
            mean_intensity = np.mean(cleaned)
            if mean_intensity < 150:  # Only for darker images
                adaptive = cv2.adaptiveThreshold(cleaned, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                                               cv2.THRESH_BINARY, 11, 2)
                # Blend with original
                cleaned = cv2.addWeighted(cleaned, 0.7, adaptive, 0.3, 0)
            
            # Convert back to RGB if needed
            if len(img.shape) == 3:
                enhanced = cv2.cvtColor(cleaned, cv2.COLOR_GRAY2RGB)
            else:
                enhanced = cleaned
            
            return enhanced
            
        except Exception:
            return img

class BalancedTextFilter:
    """
    Bộ lọc text CÂN BẰNG - Lọc text nhưng vẫn giữ được figures
    """
    
    def __init__(self):
        # Ngưỡng cân bằng - không quá nghiêm ngặt
        self.text_density_threshold = 0.7
        self.min_visual_complexity = 0.2
        self.min_diagram_score = 0.1
        self.min_figure_quality = 0.15
        
        # Thông số phân tích text nâng cao
        self.line_density_threshold = 0.25
        self.char_pattern_threshold = 0.8
        self.horizontal_structure_threshold = 0.8
        self.whitespace_ratio_threshold = 0.45
        
        # Aspect ratio filtering
        self.text_aspect_ratio_min = 0.1
        self.text_aspect_ratio_max = 12.0
        
        # Size filtering
        self.min_meaningful_size = 1000
        self.max_text_block_size = 0.75
        
        # Advanced pattern detection
        self.enable_ocr_simulation = True
        self.enable_histogram_analysis = True
        self.enable_structure_analysis = True
        
        # Debug mode
        self.debug_mode = False
        
    def analyze_and_filter_balanced(self, image_bytes, candidates):
        """
        Phân tích và lọc với độ cân bằng tốt hơn
        """
        if not CV2_AVAILABLE:
            return candidates
        
        try:
            # Validate inputs
            if not image_bytes or not candidates:
                return candidates
                
            # Đọc ảnh
            img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            img = np.array(img_pil)
            h, w = img.shape[:2]
            
            if h == 0 or w == 0:
                return candidates
            
            if self.debug_mode:
                st.write(f"🔍 Balanced Text Filter analyzing {len(candidates)} candidates")
            
            # Phân tích từng candidate với error handling
            analyzed_candidates = []
            for i, candidate in enumerate(candidates):
                try:
                    analysis = self._balanced_analyze_candidate(img, candidate)
                    candidate.update(analysis)
                    analyzed_candidates.append(candidate)
                    
                    if self.debug_mode:
                        st.write(f"   {i+1}. {candidate.get('bbox', 'N/A')}: text_score={analysis.get('text_score', 0):.2f}, is_text={analysis.get('is_text', False)}")
                except Exception as e:
                    if self.debug_mode:
                        st.warning(f"Error analyzing candidate {i+1}: {str(e)}")
                    # Keep original candidate if analysis fails
                    analyzed_candidates.append(candidate)
            
            # Lọc cân bằng
            filtered_candidates = self._balanced_filter(analyzed_candidates)
            
            if self.debug_mode:
                st.write(f"📊 Balanced filter result: {len(filtered_candidates)}/{len(candidates)}")
            
            return filtered_candidates
            
        except Exception as e:
            if self.debug_mode:
                st.error(f"❌ Balanced filter error: {str(e)}")
            return candidates  # Fallback
    
    def _balanced_analyze_candidate(self, img, candidate):
        """
        Phân tích cân bằng từng candidate với error handling
        """
        try:
            x, y, w, h = candidate['bbox']
            
            # Validate bbox
            img_h, img_w = img.shape[:2]
            if x < 0 or y < 0 or x + w > img_w or y + h > img_h or w <= 0 or h <= 0:
                return {'is_text': False, 'text_score': 0.0}
            
            roi = img[y:y+h, x:x+w]
            
            if roi.size == 0 or roi.shape[0] == 0 or roi.shape[1] == 0:
                return {'is_text': False, 'text_score': 0.0}
            
            # Các phương pháp phân tích với try-catch
            text_density = self._safe_calculate_advanced_text_density(roi)
            line_density = self._safe_analyze_line_structure(roi)
            char_pattern = self._safe_detect_character_patterns(roi)
            histogram_score = self._safe_analyze_histogram_for_text(roi)
            geometric_score = self._safe_analyze_geometric_structure(roi)
            whitespace_ratio = self._safe_calculate_whitespace_ratio(roi)
            ocr_score = self._safe_simulate_ocr_detection(roi)
            
            # Tính text score tổng hợp
            text_score = (
                text_density * 0.25 +
                line_density * 0.2 +
                char_pattern * 0.15 +
                histogram_score * 0.15 +
                ocr_score * 0.15 +
                whitespace_ratio * 0.1
            )
            
            # Aspect ratio analysis
            aspect_ratio = w / max(h, 1)  # Avoid division by zero
            is_text_aspect = (self.text_aspect_ratio_min <= aspect_ratio <= self.text_aspect_ratio_max)
            
            # Size analysis
            area = w * h
            is_text_size = area < self.min_meaningful_size
            
            # Final decision - CÂN BẰNG HỢP LÝ
            strong_text_indicators = 0
            if text_score > 0.75:
                strong_text_indicators += 1
            if line_density > 0.3:
                strong_text_indicators += 1
            if char_pattern > 0.85:
                strong_text_indicators += 1
            if whitespace_ratio > 0.5:
                strong_text_indicators += 1
            if is_text_aspect and text_score > 0.6:
                strong_text_indicators += 1
            
            # Chỉ coi là text khi có ÍT NHẤT 3 indicators mạnh
            is_text = strong_text_indicators >= 3
            
            return {
                'text_density': text_density,
                'line_density': line_density,
                'char_pattern': char_pattern,
                'histogram_score': histogram_score,
                'geometric_score': geometric_score,
                'whitespace_ratio': whitespace_ratio,
                'ocr_score': ocr_score,
                'text_score': text_score,
                'aspect_ratio': aspect_ratio,
                'is_text': is_text,
                'area': area,
                'strong_text_indicators': strong_text_indicators
            }
            
        except Exception as e:
            if self.debug_mode:
                st.warning(f"Error in candidate analysis: {str(e)}")
            return {'is_text': False, 'text_score': 0.0}
    
    def _safe_calculate_advanced_text_density(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            # Morphological text detection
            text_kernel_h = cv2.getStructuringElement(cv2.MORPH_RECT, (max(1, gray.shape[1]//10), 1))
            text_kernel_v = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(1, gray.shape[0]//10)))
            
            text_h = cv2.morphologyEx(gray, cv2.MORPH_OPEN, text_kernel_h)
            text_v = cv2.morphologyEx(gray, cv2.MORPH_OPEN, text_kernel_v)
            
            text_regions = cv2.bitwise_or(text_h, text_v)
            text_pixels = np.sum(text_regions > 0)
            total_pixels = gray.shape[0] * gray.shape[1]
            
            morphological_density = text_pixels / max(total_pixels, 1)
            
            # Edge-based text detection
            edges = cv2.Canny(gray, 50, 150)
            horizontal_edges = cv2.morphologyEx(edges, cv2.MORPH_OPEN, text_kernel_h)
            edge_density = np.sum(horizontal_edges > 0) / max(total_pixels, 1)
            
            return max(morphological_density, edge_density)
            
        except Exception:
            return 0.0
    
    def _safe_analyze_line_structure(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(1, gray.shape[1]//5), 1))
            horizontal_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, horizontal_kernel)
            
            contours, _ = cv2.findContours(horizontal_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            line_count = len(contours)
            
            height = gray.shape[0]
            line_density = line_count / max(height / 20, 1)
            
            return min(1.0, line_density)
            
        except Exception:
            return 0.0
    
    def _safe_detect_character_patterns(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            binary = cv2.bitwise_not(binary)
            
            num_labels, labels, stats, centroids = cv2.connectedComponentsWithStats(binary)
            
            char_like_components = 0
            total_area = gray.shape[0] * gray.shape[1]
            
            for i in range(1, min(num_labels, 100)):  # Limit to avoid memory issues
                area = stats[i, cv2.CC_STAT_AREA]
                width = stats[i, cv2.CC_STAT_WIDTH]
                height = stats[i, cv2.CC_STAT_HEIGHT]
                
                if (50 < area < 1000 and
                    5 < width < 50 and
                    10 < height < 50 and
                    0.2 < width/max(height, 1) < 3.0):
                    char_like_components += 1
            
            char_density = char_like_components / max(total_area / 500, 1)
            return min(1.0, char_density)
            
        except Exception:
            return 0.0
    
    def _safe_analyze_histogram_for_text(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            hist = cv2.calcHist([gray], [0], None, [256], [0, 256])
            hist = hist.flatten()
            
            # Find peaks
            peaks = []
            for i in range(1, len(hist) - 1):
                if hist[i] > hist[i-1] and hist[i] > hist[i+1] and hist[i] > np.max(hist) * 0.1:
                    peaks.append(i)
            
            if len(peaks) >= 2:
                peak_distances = [abs(peaks[i+1] - peaks[i]) for i in range(len(peaks) - 1)]
                if max(peak_distances) > 100:
                    return 0.8
            
            # Calculate entropy
            hist_norm = hist / max(np.sum(hist), 1)
            entropy = -np.sum(hist_norm * np.log2(hist_norm + 1e-10))
            
            if entropy < 4.0:
                return 0.6
            
            return 0.2
            
        except Exception:
            return 0.0
    
    def _safe_analyze_geometric_structure(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            edges = cv2.Canny(gray, 50, 150)
            
            # Detect lines
            lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=30, minLineLength=20, maxLineGap=10)
            line_count = len(lines) if lines is not None else 0
            
            # Detect circles
            circles = cv2.HoughCircles(gray, cv2.HOUGH_GRADIENT, dp=1, minDist=20, 
                                     param1=50, param2=30, minRadius=5, maxRadius=100)
            circle_count = len(circles[0]) if circles is not None else 0
            
            # Detect complex contours
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            complex_contours = 0
            
            for contour in contours[:20]:  # Limit processing
                area = cv2.contourArea(contour)
                if area > 500:
                    hull = cv2.convexHull(contour)
                    hull_area = cv2.contourArea(hull)
                    if hull_area > 0:
                        solidity = area / hull_area
                        if solidity < 0.8:
                            complex_contours += 1
            
            total_area = gray.shape[0] * gray.shape[1]
            geometric_score = (line_count * 0.1 + circle_count * 0.5 + complex_contours * 0.3) / max(total_area / 1000, 1)
            
            return min(1.0, geometric_score)
            
        except Exception:
            return 0.0
    
    def _safe_calculate_whitespace_ratio(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            _, binary = cv2.threshold(gray, 200, 255, cv2.THRESH_BINARY)
            
            white_pixels = np.sum(binary == 255)
            total_pixels = gray.shape[0] * gray.shape[1]
            
            return white_pixels / max(total_pixels, 1)
            
        except Exception:
            return 0.0
    
    def _safe_simulate_ocr_detection(self, roi):
        """Safe version with error handling"""
        try:
            gray = cv2.cvtColor(roi, cv2.COLOR_RGB2GRAY) if len(roi.shape) == 3 else roi
            
            if gray.shape[0] == 0 or gray.shape[1] == 0:
                return 0.0
            
            # Resize to standard height
            target_height = 32
            scale = target_height / max(gray.shape[0], 1)
            new_width = max(1, int(gray.shape[1] * scale))
            
            resized = cv2.resize(gray, (new_width, target_height))
            enhanced = cv2.equalizeHist(resized)
            
            # Horizontal projections
            h_projection = np.sum(enhanced < 128, axis=1)
            
            # Count peaks
            h_peaks = 0
            for i in range(1, len(h_projection) - 1):
                if h_projection[i] > h_projection[i-1] and h_projection[i] > h_projection[i+1]:
                    if h_projection[i] > np.max(h_projection) * 0.3:
                        h_peaks += 1
            
            if h_peaks >= 2:
                return 0.9
            elif h_peaks == 1:
                return 0.7
            else:
                return 0.3
                
        except Exception:
            return 0.0
    
    def _balanced_filter(self, candidates):
        """
        Lọc cân bằng - ưu tiên giữ lại figures
        """
        filtered = []
        
        for candidate in candidates:
            try:
                # Chỉ loại bỏ khi RẤT CHẮC CHẮN là text
                if candidate.get('is_text', False):
                    # Cho phép giữ lại nếu có geometric complexity cao
                    geometric_score = candidate.get('geometric_score', 0)
                    if geometric_score >= 0.3:
                        candidate['override_reason'] = 'complex_geometry'
                        filtered.append(candidate)
                        continue
                    
                    # Cho phép giữ lại nếu kích thước lớn và có structure
                    area = candidate.get('area', 0)
                    if area > 5000 and geometric_score > 0.1:
                        candidate['override_reason'] = 'large_with_structure'
                        filtered.append(candidate)
                        continue
                    
                    # Loại bỏ text chắc chắn
                    continue
                
                # Kiểm tra các điều kiện khác
                text_score = candidate.get('text_score', 0)
                if text_score > self.text_density_threshold:
                    geometric_score = candidate.get('geometric_score', 0)
                    if geometric_score >= self.min_diagram_score:
                        candidate['override_reason'] = 'has_diagram_elements'
                        filtered.append(candidate)
                    continue
                
                # Kiểm tra size
                area = candidate.get('area', 0)
                if area < self.min_meaningful_size:
                    geometric_score = candidate.get('geometric_score', 0)
                    if geometric_score >= 0.4:
                        candidate['override_reason'] = 'small_but_complex'
                        filtered.append(candidate)
                    continue
                
                # Nếu pass hầu hết tests thì giữ lại
                filtered.append(candidate)
                
            except Exception as e:
                # If error in filtering, keep the candidate
                if self.debug_mode:
                    st.warning(f"Error filtering candidate: {str(e)}")
                filtered.append(candidate)
        
        return filtered

class EnhancedContentBasedFigureFilter:
    """
    Bộ lọc thông minh với Mistral OCR Integration
    """
    
    def __init__(self, mistral_ocr_service=None):
        self.text_filter = BalancedTextFilter()
        self.enable_balanced_filter = True
        self.min_estimated_count = 1
        self.max_estimated_count = 15
        self.mistral_ocr = mistral_ocr_service
        self.enable_ocr_counting = True
        
    def analyze_content_and_filter_with_ocr(self, image_bytes, candidates):
        """
        Phân tích với Mistral OCR + Balanced Text Filter
        """
        if not CV2_AVAILABLE:
            return candidates
        
        try:
            # OCR Analysis để đếm figures
            estimated_count = self.min_estimated_count
            ocr_info = {}
            
            if self.mistral_ocr and self.enable_ocr_counting:
                with st.spinner("🤖 Analyzing image content with Mistral OCR..."):
                    ocr_result = self.mistral_ocr.analyze_image_content(image_bytes)
                    
                    if ocr_result['success']:
                        estimated_count = max(ocr_result['total_count'], self.min_estimated_count)
                        estimated_count = min(estimated_count, self.max_estimated_count)
                        ocr_info = ocr_result
                        
                        # Display Mistral OCR results with special styling
                        st.markdown(f"""
                        <div class="mistral-badge">
                            🧠 Mistral OCR: {ocr_result['figure_count']} figures, {ocr_result['table_count']} tables 
                            (confidence: {ocr_result['confidence']:.1f}, model: {ocr_result.get('model', 'pixtral-12b')})
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Show analysis details if available
                        if ocr_result.get('analysis'):
                            analysis = ocr_result['analysis']
                            if analysis.get('visual_complexity'):
                                st.info(f"📊 Visual complexity: {analysis.get('visual_complexity')} | Content type: {analysis.get('content_type', 'unknown')}")
                        
                    else:
                        # Fallback to conservative estimation
                        img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                        img = np.array(img_pil)
                        estimated_count = self._estimate_figure_count_conservative(img)
                        st.info(f"📊 Conservative estimate: {estimated_count} figures (Mistral OCR fallback)")
            else:
                # Original estimation method
                img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                img = np.array(img_pil)
                estimated_count = self._estimate_figure_count_conservative(img)
                st.info(f"📊 Estimated: {estimated_count} figures (traditional method)")
            
            # Balanced Text Filter
            if self.enable_balanced_filter:
                filtered_candidates = self.text_filter.analyze_and_filter_balanced(image_bytes, candidates)
                st.success(f"🧠 Balanced Text Filter: {len(filtered_candidates)}/{len(candidates)} figures → target: {estimated_count}")
            else:
                filtered_candidates = candidates
            
            # Intelligent filtering based on OCR results
            if ocr_info.get('success') and ocr_info.get('visual_elements'):
                # Use Mistral visual elements to improve filtering
                filtered_candidates = self._filter_with_mistral_analysis(filtered_candidates, ocr_info)
            
            # Adjust count based on estimation
            target_count = min(estimated_count + 1, self.max_estimated_count)  # +1 buffer
            if len(filtered_candidates) > target_count:
                # Sort by confidence and take top candidates
                sorted_candidates = sorted(filtered_candidates, 
                                         key=lambda x: x.get('final_confidence', 0), reverse=True)
                filtered_candidates = sorted_candidates[:target_count]
                st.info(f"🎯 Limited to top {target_count} figures based on Mistral OCR estimate")
            
            return filtered_candidates
            
        except Exception as e:
            st.error(f"❌ Enhanced filter error: {str(e)}")
            return candidates
    
    def _filter_with_mistral_analysis(self, candidates, ocr_info):
        """
        Sử dụng Mistral analysis để cải thiện filtering
        """
        try:
            visual_elements = ocr_info.get('visual_elements', [])
            analysis = ocr_info.get('analysis', {})
            
            if not visual_elements:
                return candidates
            
            enhanced_candidates = []
            
            for candidate in candidates:
                bbox = candidate['bbox']
                
                # Boost confidence based on visual complexity analysis
                if analysis.get('visual_complexity') == 'high':
                    candidate['final_confidence'] = candidate.get('final_confidence', 50) + 10
                    candidate['mistral_boost'] = 'high_complexity'
                
                # Boost confidence for mathematical content
                if analysis.get('has_mathematical_content'):
                    candidate['final_confidence'] = candidate.get('final_confidence', 50) + 15
                    candidate['mistral_boost'] = 'mathematical_content'
                
                # Boost confidence based on content type
                content_type = analysis.get('content_type', '')
                if content_type in ['exam', 'textbook']:
                    candidate['final_confidence'] = candidate.get('final_confidence', 50) + 8
                    candidate['mistral_boost'] = f'content_type_{content_type}'
                
                enhanced_candidates.append(candidate)
            
            return enhanced_candidates
            
        except Exception:
            return candidates
    
    def _estimate_figure_count_conservative(self, img):
        """
        Ước tính conservative số lượng figures (fallback method)
        """
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            h, w = gray.shape
            
            # Phân tích layout đơn giản
            # Detect horizontal separators
            h_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (w//20, 1))
            h_lines = cv2.morphologyEx(gray, cv2.MORPH_OPEN, h_kernel)
            h_separators = len(cv2.findContours(h_lines, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)[0])
            
            # Estimate based on separators
            estimated = min(max(h_separators + 1, self.min_estimated_count), self.max_estimated_count)
            
            return estimated
            
        except Exception:
            return 3  # Safe fallback

class GeminiAPI:
    def __init__(self, api_key: str):
        self.api_key = api_key
        self.base_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent"
        self.session = requests.Session()
        self.max_retries = 3
        self.timeout = 120
    
    def encode_image(self, image_data: bytes) -> str:
        return base64.b64encode(image_data).decode('utf-8')
    
    def convert_to_latex(self, content_data: bytes, content_type: str, prompt: str) -> str:
        headers = {"Content-Type": "application/json"}
        
        if content_type.startswith('image/'):
            mime_type = content_type
        else:
            mime_type = "image/png"
        
        # Check image size
        if len(content_data) > 20 * 1024 * 1024:  # 20MB limit
            raise Exception("Image quá lớn (>20MB). Vui lòng resize ảnh.")
        
        encoded_content = self.encode_image(content_data)
        
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": prompt},
                        {
                            "inline_data": {
                                "mime_type": mime_type,
                                "data": encoded_content
                            }
                        }
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.1,
                "topK": 1,
                "topP": 0.8,
                "maxOutputTokens": 8192,
            }
        }
        
        for attempt in range(self.max_retries):
            try:
                response = self.session.post(
                    f"{self.base_url}?key={self.api_key}",
                    headers=headers,
                    json=payload,
                    timeout=self.timeout
                )
                
                if response.status_code == 200:
                    result = response.json()
                    if 'candidates' in result and len(result['candidates']) > 0:
                        content = result['candidates'][0]['content']['parts'][0]['text']
                        return content.strip()
                    else:
                        raise Exception("API không trả về kết quả hợp lệ")
                elif response.status_code == 401:
                    raise Exception("API key không hợp lệ hoặc đã hết hạn")
                elif response.status_code == 429:
                    if attempt < self.max_retries - 1:
                        time.sleep(2 ** attempt)  # Exponential backoff
                        continue
                    raise Exception("Đã vượt quá giới hạn rate limit")
                else:
                    error_text = response.text[:200] if response.text else "Unknown error"
                    raise Exception(f"API Error {response.status_code}: {error_text}")
            
            except requests.exceptions.Timeout:
                if attempt < self.max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
                raise Exception("Request timeout - thử lại sau ít phút")
            except requests.exceptions.ConnectionError:
                if attempt < self.max_retries - 1:
                    time.sleep(2 ** attempt)
                    continue
                raise Exception("Lỗi kết nối mạng")
            except Exception as e:
                if attempt < self.max_retries - 1 and "rate limit" in str(e).lower():
                    time.sleep(2 ** attempt)
                    continue
                raise Exception(str(e))

class PDFProcessor:
    @staticmethod
    def extract_images_and_text(pdf_file, max_pages=None):
        """Extract images with memory management"""
        try:
            # Read file content
            file_content = pdf_file.read()
            if len(file_content) == 0:
                raise Exception("PDF file is empty")
            
            pdf_document = fitz.open(stream=file_content, filetype="pdf")
            images = []
            
            total_pages = pdf_document.page_count
            if max_pages:
                total_pages = min(total_pages, max_pages)
            
            for page_num in range(total_pages):
                try:
                    page = pdf_document[page_num]
                    
                    # Use reasonable resolution to avoid memory issues
                    mat = fitz.Matrix(2.0, 2.0)  # Reduced from 3.5 to 2.0
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    
                    # Clean up pixmap
                    pix = None
                    
                    img = Image.open(io.BytesIO(img_data))
                    
                    # Limit image size to prevent memory issues
                    max_size = (2000, 2000)
                    if img.size[0] > max_size[0] or img.size[1] > max_size[1]:
                        img.thumbnail(max_size, Image.Resampling.LANCZOS)
                    
                    images.append((img, page_num + 1))
                    
                    # Force garbage collection every few pages
                    if page_num % 5 == 0:
                        gc.collect()
                        
                except Exception as e:
                    st.warning(f"Lỗi xử lý trang {page_num + 1}: {str(e)}")
                    continue
            
            pdf_document.close()
            return images
            
        except Exception as e:
            raise Exception(f"Lỗi đọc PDF: {str(e)}")

class SuperEnhancedImageExtractor:
    """
    Tách ảnh với Balanced Text Filter + Mistral OCR Integration
    """
    
    def __init__(self, mistral_ocr_service=None):
        # Tham số cơ bản
        self.min_area_ratio = 0.0005
        self.min_area_abs = 400
        self.min_width = 20
        self.min_height = 20
        self.max_figures = 25
        self.max_area_ratio = 0.80
        
        # Tham số cắt ảnh
        self.smart_padding = 30
        self.quality_threshold = 0.15
        self.edge_margin = 0.005
        
        # Tham số confidence
        self.confidence_threshold = 15
        self.final_confidence_threshold = 65
        
        # Tham số morphology
        self.morph_kernel_size = 2
        self.dilate_iterations = 1
        self.erode_iterations = 1
        
        # Tham số edge detection
        self.canny_low = 30
        self.canny_high = 80
        self.blur_kernel = 3
        
        # Enhanced Content-Based Filter với Mistral OCR
        self.content_filter = EnhancedContentBasedFigureFilter(mistral_ocr_service)
        self.enable_content_filter = True
        
        # Debug mode
        self.debug_mode = False
    
    def extract_figures_and_tables(self, image_bytes, start_img_idx=0, start_table_idx=0):
        """
        Tách ảnh với Balanced Text Filter và continuous numbering
        """
        if not CV2_AVAILABLE:
            return [], 0, 0, start_img_idx, start_table_idx
        
        try:
            # Validate input
            if not image_bytes or len(image_bytes) == 0:
                return [], 0, 0, start_img_idx, start_table_idx
            
            # Đọc ảnh với error handling
            try:
                img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                
                # Limit image size to prevent memory issues
                max_size = (3000, 3000)
                if img_pil.size[0] > max_size[0] or img_pil.size[1] > max_size[1]:
                    img_pil.thumbnail(max_size, Image.Resampling.LANCZOS)
                
                img = np.array(img_pil)
                h, w = img.shape[:2]
                
                if h == 0 or w == 0:
                    return [], 0, 0, start_img_idx, start_table_idx
                    
            except Exception as e:
                st.error(f"Lỗi đọc ảnh: {str(e)}")
                return [], 0, 0, start_img_idx, start_table_idx
            
            # Tiền xử lý
            enhanced_img = self._enhance_image(img)
            
            # Tách ảnh bằng 4 phương pháp với error handling
            all_candidates = []
            
            try:
                edge_candidates = self._detect_by_edges(enhanced_img, w, h)
                all_candidates.extend(edge_candidates)
            except Exception as e:
                if self.debug_mode:
                    st.warning(f"Edge detection error: {str(e)}")
            
            try:
                contour_candidates = self._detect_by_contours(enhanced_img, w, h)
                all_candidates.extend(contour_candidates)
            except Exception as e:
                if self.debug_mode:
                    st.warning(f"Contour detection error: {str(e)}")
            
            try:
                grid_candidates = self._detect_by_grid(enhanced_img, w, h)
                all_candidates.extend(grid_candidates)
            except Exception as e:
                if self.debug_mode:
                    st.warning(f"Grid detection error: {str(e)}")
            
            try:
                blob_candidates = self._detect_by_blobs(enhanced_img, w, h)
                all_candidates.extend(blob_candidates)
            except Exception as e:
                if self.debug_mode:
                    st.warning(f"Blob detection error: {str(e)}")
            
            # Lọc và merge
            filtered_candidates = self._filter_and_merge_candidates(all_candidates, w, h)
            
            # Enhanced Content-Based Filter với Mistral OCR
            if self.enable_content_filter:
                try:
                    content_filtered = self.content_filter.analyze_content_and_filter_with_ocr(image_bytes, filtered_candidates)
                    filtered_candidates = content_filtered
                except Exception as e:
                    if self.debug_mode:
                        st.warning(f"Content filter error: {str(e)}")
            
            # Tạo final figures với continuous numbering
            final_figures, final_img_idx, final_table_idx = self._create_final_figures(
                filtered_candidates, img, w, h, start_img_idx, start_table_idx
            )
            
            return final_figures, h, w, final_img_idx, final_table_idx
            
        except Exception as e:
            st.error(f"❌ Extraction error: {str(e)}")
            return [], 0, 0, start_img_idx, start_table_idx
    
    def _enhance_image(self, img):
        """
        Tiền xử lý ảnh với error handling
        """
        try:
            gray = cv2.cvtColor(img, cv2.COLOR_RGB2GRAY)
            blurred = cv2.GaussianBlur(gray, (self.blur_kernel, self.blur_kernel), 0)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(blurred)
            return cv2.normalize(enhanced, None, 0, 255, cv2.NORM_MINMAX)
        except Exception:
            # Fallback to simple grayscale
            return cv2.cvtColor(img, cv2.COLOR_RGB2GRAY) if len(img.shape) == 3 else img
    
    def _detect_by_edges(self, gray_img, w, h):
        """Edge detection với error handling"""
        try:
            edges = cv2.Canny(gray_img, self.canny_low, self.canny_high)
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (3, 3))
            edges_dilated = cv2.dilate(edges, kernel, iterations=1)
            
            contours, _ = cv2.findContours(edges_dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            candidates = []
            for cnt in contours:
                try:
                    x, y, ww, hh = cv2.boundingRect(cnt)
                    area = ww * hh
                    
                    if self._is_valid_candidate(x, y, ww, hh, area, w, h):
                        candidates.append({
                            'bbox': (x, y, ww, hh),
                            'area': area,
                            'method': 'edge',
                            'confidence': 25
                        })
                except Exception:
                    continue
            
            return candidates
        except Exception:
            return []
    
    def _detect_by_contours(self, gray_img, w, h):
        """Contour detection với error handling"""
        try:
            _, binary = cv2.threshold(gray_img, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (self.morph_kernel_size, self.morph_kernel_size))
            binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
            contours, _ = cv2.findContours(binary, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            candidates = []
            for cnt in contours:
                try:
                    x, y, ww, hh = cv2.boundingRect(cnt)
                    area = ww * hh
                    
                    if self._is_valid_candidate(x, y, ww, hh, area, w, h):
                        candidates.append({
                            'bbox': (x, y, ww, hh),
                            'area': area,
                            'method': 'contour',
                            'confidence': 30
                        })
                except Exception:
                    continue
            
            return candidates
        except Exception:
            return []
    
    def _detect_by_grid(self, gray_img, w, h):
        """Grid detection với error handling"""
        try:
            horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (max(1, w//20), 1))
            horizontal_lines = cv2.morphologyEx(gray_img, cv2.MORPH_OPEN, horizontal_kernel)
            
            vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, max(1, h//20)))
            vertical_lines = cv2.morphologyEx(gray_img, cv2.MORPH_OPEN, vertical_kernel)
            
            grid_mask = cv2.bitwise_or(horizontal_lines, vertical_lines)
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (5, 5))
            grid_dilated = cv2.dilate(grid_mask, kernel, iterations=2)
            
            contours, _ = cv2.findContours(grid_dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            candidates = []
            for cnt in contours:
                try:
                    x, y, ww, hh = cv2.boundingRect(cnt)
                    area = ww * hh
                    
                    if self._is_valid_candidate(x, y, ww, hh, area, w, h):
                        aspect_ratio = ww / max(hh, 1)
                        confidence = 50 if aspect_ratio > 1.5 else 30
                        
                        candidates.append({
                            'bbox': (x, y, ww, hh),
                            'area': area,
                            'method': 'grid',
                            'confidence': confidence,
                            'is_table': aspect_ratio > 1.5
                        })
                except Exception:
                    continue
            
            return candidates
        except Exception:
            return []
    
    def _detect_by_blobs(self, gray_img, w, h):
        """Blob detection với error handling"""
        try:
            adaptive_thresh = cv2.adaptiveThreshold(
                gray_img, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2
            )
            
            inverted = cv2.bitwise_not(adaptive_thresh)
            kernel = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (5, 5))
            opened = cv2.morphologyEx(inverted, cv2.MORPH_OPEN, kernel)
            
            contours, _ = cv2.findContours(opened, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            candidates = []
            for cnt in contours:
                try:
                    x, y, ww, hh = cv2.boundingRect(cnt)
                    area = ww * hh
                    
                    if self._is_valid_candidate(x, y, ww, hh, area, w, h):
                        candidates.append({
                            'bbox': (x, y, ww, hh),
                            'area': area,
                            'method': 'blob',
                            'confidence': 28
                        })
                except Exception:
                    continue
            
            return candidates
        except Exception:
            return []
    
    def _is_valid_candidate(self, x, y, ww, hh, area, img_w, img_h):
        """
        Kiểm tra candidate có hợp lệ với better validation
        """
        try:
            # Basic validation
            if x < 0 or y < 0 or ww <= 0 or hh <= 0:
                return False
            
            if x + ww > img_w or y + hh > img_h:
                return False
            
            area_ratio = area / max(img_w * img_h, 1)
            
            if (area < self.min_area_abs or 
                area_ratio < self.min_area_ratio or 
                area_ratio > self.max_area_ratio or
                ww < self.min_width or 
                hh < self.min_height):
                return False
            
            if (x < self.edge_margin * img_w or 
                y < self.edge_margin * img_h or 
                (x + ww) > (1 - self.edge_margin) * img_w or 
                (y + hh) > (1 - self.edge_margin) * img_h):
                return False
            
            return True
        except Exception:
            return False
    
    def _filter_and_merge_candidates(self, candidates, w, h):
        """
        Lọc và merge candidates với error handling
        """
        try:
            if not candidates:
                return []
            
            candidates = sorted(candidates, key=lambda x: x.get('area', 0), reverse=True)
            
            filtered = []
            for candidate in candidates:
                try:
                    if not self._is_overlapping_with_list(candidate, filtered):
                        candidate['final_confidence'] = self._calculate_final_confidence(candidate, w, h)
                        if candidate['final_confidence'] >= self.confidence_threshold:
                            filtered.append(candidate)
                except Exception:
                    continue
            
            return filtered[:self.max_figures]
        except Exception:
            return []
    
    def _is_overlapping_with_list(self, candidate, existing_list):
        """
        Kiểm tra overlap với error handling
        """
        try:
            x1, y1, w1, h1 = candidate['bbox']
            
            for existing in existing_list:
                x2, y2, w2, h2 = existing['bbox']
                
                intersection_area = max(0, min(x1+w1, x2+w2) - max(x1, x2)) * max(0, min(y1+h1, y2+h2) - max(y1, y2))
                union_area = w1*h1 + w2*h2 - intersection_area
                
                if union_area > 0:
                    iou = intersection_area / union_area
                    if iou > 0.25:
                        return True
            
            return False
        except Exception:
            return False
    
    def _calculate_final_confidence(self, candidate, w, h):
        """
        Tính confidence với error handling
        """
        try:
            x, y, ww, hh = candidate['bbox']
            area_ratio = candidate['area'] / max(w * h, 1)
            aspect_ratio = ww / max(hh, 1)
            
            confidence = candidate.get('confidence', 20)
            
            # Bonus cho size phù hợp
            if 0.015 < area_ratio < 0.5:
                confidence += 20
            elif 0.005 < area_ratio < 0.7:
                confidence += 10
            
            # Bonus cho aspect ratio
            if 0.4 < aspect_ratio < 4.0:
                confidence += 15
            elif 0.2 < aspect_ratio < 6.0:
                confidence += 8
            
            # Bonus cho method
            if candidate['method'] == 'grid':
                confidence += 12
            elif candidate['method'] == 'edge':
                confidence += 8
            
            return min(100, confidence)
        except Exception:
            return 20
    
    def _create_final_figures(self, candidates, img, w, h, start_img_idx=0, start_table_idx=0):
        """
        Tạo final figures với confidence filter và continuous numbering
        """
        try:
            candidates = sorted(candidates, key=lambda x: (x['bbox'][1], x['bbox'][0]))
            
            # Lọc theo final confidence threshold
            high_confidence_candidates = [c for c in candidates 
                                        if c.get('final_confidence', 0) >= self.final_confidence_threshold]
            
            if self.debug_mode:
                st.write(f"🎯 Confidence Filter: {len(high_confidence_candidates)}/{len(candidates)} figures above {self.final_confidence_threshold}%")
            elif len(candidates) > 0:
                st.info(f"🎯 Confidence Filter: Giữ {len(high_confidence_candidates)}/{len(candidates)} figures có confidence ≥{self.final_confidence_threshold}%")
            
            final_figures = []
            img_idx = start_img_idx
            table_idx = start_table_idx
            
            for candidate in high_confidence_candidates:
                try:
                    cropped_img = self._smart_crop(img, candidate, w, h)
                    
                    if cropped_img is None:
                        continue
                    
                    buf = io.BytesIO()
                    Image.fromarray(cropped_img).save(buf, format="JPEG", quality=95)
                    b64 = base64.b64encode(buf.getvalue()).decode()
                    
                    is_table = candidate.get('is_table', False) or candidate.get('method') == 'grid'
                    
                    if is_table:
                        table_idx += 1
                        name = f"table-{table_idx}.jpeg"
                    else:
                        img_idx += 1
                        name = f"figure-{img_idx}.jpeg"
                    
                    final_figures.append({
                        "name": name,
                        "base64": b64,
                        "is_table": is_table,
                        "bbox": candidate["bbox"],
                        "confidence": candidate["final_confidence"],
                        "area_ratio": candidate["area"] / max(w * h, 1),
                        "aspect_ratio": candidate["bbox"][2] / max(candidate["bbox"][3], 1),
                        "method": candidate["method"],
                        "center_y": candidate["bbox"][1] + candidate["bbox"][3] // 2,
                        "center_x": candidate["bbox"][0] + candidate["bbox"][2] // 2,
                        "override_reason": candidate.get("override_reason", None),
                        "mistral_boost": candidate.get("mistral_boost", None)
                    })
                except Exception as e:
                    if self.debug_mode:
                        st.warning(f"Error creating figure: {str(e)}")
                    continue
            
            return final_figures, img_idx, table_idx
        except Exception:
            return [], start_img_idx, start_table_idx
    
    def _smart_crop(self, img, candidate, img_w, img_h):
        """
        Cắt ảnh thông minh với error handling
        """
        try:
            x, y, w, h = candidate['bbox']
            
            # Validate bounds
            if x < 0 or y < 0 or x + w > img_w or y + h > img_h or w <= 0 or h <= 0:
                return None
            
            padding_x = min(self.smart_padding, w // 4)
            padding_y = min(self.smart_padding, h // 4)
            
            x0 = max(0, x - padding_x)
            y0 = max(0, y - padding_y)
            x1 = min(img_w, x + w + padding_x)
            y1 = min(img_h, y + h + padding_y)
            
            cropped = img[y0:y1, x0:x1]
            
            if cropped.size == 0 or cropped.shape[0] == 0 or cropped.shape[1] == 0:
                return None
            
            return cropped
        except Exception:
            return None
    
    def insert_figures_into_text_precisely(self, text, figures, img_h, img_w, show_override_info=True):
        """
        Chèn figures vào text với option hiển thị override info
        """
        try:
            if not figures:
                return text
            
            lines = text.split('\n')
            sorted_figures = sorted(figures, key=lambda f: f['center_y'])
            
            result_lines = lines[:]
            offset = 0
            
            for i, figure in enumerate(sorted_figures):
                try:
                    insertion_line = self._calculate_insertion_position(figure, lines, i, len(sorted_figures))
                    actual_insertion = insertion_line + offset
                    
                    if actual_insertion > len(result_lines):
                        actual_insertion = len(result_lines)
                    
                    if figure['is_table']:
                        tag = f"[📊 BẢNG: {figure['name']}]"
                    else:
                        tag = f"[🖼️ HÌNH: {figure['name']}]"
                    
                    # Thêm thông tin override nếu có và được yêu cầu
                    if show_override_info and figure.get('override_reason'):
                        tag += f" (kept: {figure['override_reason']})"
                    
                    # Thêm thông tin Mistral boost nếu có
                    if show_override_info and figure.get('mistral_boost'):
                        tag += f" (🧠 {figure['mistral_boost']})"
                    
                    result_lines.insert(actual_insertion, "")
                    result_lines.insert(actual_insertion + 1, tag)
                    result_lines.insert(actual_insertion + 2, "")
                    
                    offset += 3
                except Exception:
                    continue
            
            return '\n'.join(result_lines)
        except Exception:
            return text
    
    def _calculate_insertion_position(self, figure, lines, fig_index, total_figures):
        """
        Tính vị trí chèn với error handling
        """
        try:
            question_lines = []
            for i, line in enumerate(lines):
                if re.match(r'^(câu|bài|question)\s*\d+', line.strip().lower()):
                    question_lines.append(i)
            
            if question_lines:
                if fig_index < len(question_lines):
                    return question_lines[fig_index] + 1
                else:
                    return question_lines[-1] + 2
            
            section_size = max(1, len(lines) // (total_figures + 1))
            return min(section_size * (fig_index + 1), len(lines) - 1)
        except Exception:
            return 0
    
    def create_beautiful_debug_visualization(self, image_bytes, figures):
        """
        Tạo debug visualization với error handling
        """
        try:
            img_pil = Image.open(io.BytesIO(image_bytes)).convert("RGB")
            draw = ImageDraw.Draw(img_pil)
            
            colors = ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4', '#FFEAA7', '#DDA0DD']
            
            for i, fig in enumerate(figures):
                try:
                    color = colors[i % len(colors)]
                    x, y, w, h = fig['bbox']
                    
                    draw.rectangle([x, y, x+w, y+h], outline=color, width=4)
                    
                    # Corner markers
                    corner_size = 10
                    draw.rectangle([x, y, x+corner_size, y+corner_size], fill=color)
                    draw.rectangle([x+w-corner_size, y, x+w, y+corner_size], fill=color)
                    
                    # Center point
                    center_x, center_y = fig['center_x'], fig['center_y']
                    draw.ellipse([center_x-8, center_y-8, center_x+8, center_y+8], fill=color)
                    
                    # Simple label với Mistral boost info
                    label = f"{fig['name']} ({fig['confidence']:.0f}%)"
                    if fig.get('override_reason'):
                        label += f" [{fig['override_reason']}]"
                    if fig.get('mistral_boost'):
                        label += f" 🧠{fig['mistral_boost']}"
                    draw.text((x + 5, y + 5), label, fill=color, stroke_width=2, stroke_fill='white')
                except Exception:
                    continue
            
            return img_pil
        except Exception:
            return None

class EnhancedWordExporter:
    """
    Xuất Word document với chèn figures đúng vị trí
    """
    
    @staticmethod
    def create_word_document(latex_content: str, extracted_figures=None, images=None) -> io.BytesIO:
        try:
            if not DOCX_AVAILABLE:
                raise Exception("python-docx not available")
            
            doc = Document()
            
            # Cấu hình font
            style = doc.styles['Normal']
            style.font.name = 'Times New Roman'
            style.font.size = Pt(12)
            
            # Xử lý nội dung LaTeX
            lines = latex_content.split('\n')
            
            for line in lines:
                line = line.strip()
                
                if not line or line.startswith('<!--'):
                    continue
                
                if line.startswith('```'):
                    continue
                
                # Xử lý tags hình ảnh
                if line.startswith('[') and line.endswith(']'):
                    if 'HÌNH:' in line or 'BẢNG:' in line:
                        EnhancedWordExporter._insert_figure_to_word(doc, line, extracted_figures)
                        continue
                
                # Xử lý câu hỏi - đặt màu đen và in đậm
                if re.match(r'^(câu|bài)\s+\d+', line.lower()):
                    heading = doc.add_heading(line, level=3)
                    # Đặt màu đen cho câu hỏi và in đậm
                    for run in heading.runs:
                        run.font.color.rgb = RGBColor(0, 0, 0)  # Màu đen
                        run.font.bold = True
                    continue
                
                # Xử lý paragraph thường
                if line:
                    para = doc.add_paragraph()
                    EnhancedWordExporter._process_latex_content(para, line)
            
            # Lưu vào buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer
            
        except Exception as e:
            st.error(f"❌ Lỗi tạo Word: {str(e)}")
            raise e
    
    @staticmethod
    def _process_latex_content(para, content):
        """
        Xử lý nội dung LaTeX - chuyển ${...}$ thành dạng Word hiệu quả
        """
        # Tách content thành các phần: text thường và công thức ${...}$
        parts = re.split(r'(\$\{[^}]+\}\$)', content)
        
        for part in parts:
            if part.startswith('${') and part.endswith('}
                
                # Chuyển đổi một số ký hiệu LaTeX cơ bản thành Unicode
                formula_content = EnhancedWordExporter._convert_latex_to_unicode(formula_content)
                
                # Thêm công thức vào paragraph với font khác biệt
                run = para.add_run(formula_content)
                run.font.name = 'Cambria Math'  # Font phù hợp cho toán học
                run.font.italic = True  # In nghiêng cho công thức
                
            elif part.strip():
                # Đây là text thường
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    @staticmethod
    def _convert_latex_to_unicode(latex_content):
        """
        Chuyển đổi một số ký hiệu LaTeX sang Unicode
        """
        # Dictionary chuyển đổi LaTeX sang Unicode
        latex_to_unicode = {
            # Chữ Hy Lạp
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\epsilon': 'ε', '\\theta': 'θ', '\\lambda': 'λ', '\\mu': 'μ',
            '\\pi': 'π', '\\sigma': 'σ', '\\phi': 'φ', '\\omega': 'ω',
            '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ', '\\Pi': 'Π',
            '\\Sigma': 'Σ', '\\Phi': 'Φ', '\\Omega': 'Ω',
            
            # Ký hiệu toán học
            '\\infty': '∞', '\\pm': '±', '\\mp': '∓',
            '\\times': '×', '\\div': '÷', '\\cdot': '·',
            '\\leq': '≤', '\\geq': '≥', '\\neq': '≠',
            '\\approx': '≈', '\\equiv': '≡', '\\sim': '∼',
            '\\subset': '⊂', '\\supset': '⊃', '\\in': '∈',
            '\\notin': '∉', '\\cup': '∪', '\\cap': '∩',
            '\\sum': '∑', '\\prod': '∏', '\\int': '∫',
            '\\partial': '∂', '\\nabla': '∇',
            
            # Mũi tên
            '\\rightarrow': '→', '\\leftarrow': '←',
            '\\leftrightarrow': '↔', '\\Rightarrow': '⇒',
            '\\Leftarrow': '⇐', '\\Leftrightarrow': '⇔',
            
            # Xử lý phân số đơn giản
            '\\frac{1}{2}': '½', '\\frac{1}{3}': '⅓', '\\frac{2}{3}': '⅔',
            '\\frac{1}{4}': '¼', '\\frac{3}{4}': '¾', '\\frac{1}{8}': '⅛',
            
            # Lũy thừa đơn giản (sử dụng superscript Unicode)
            '^2': '²', '^3': '³', '^1': '¹',
            '^0': '⁰', '^4': '⁴', '^5': '⁵',
            '^6': '⁶', '^7': '⁷', '^8': '⁸', '^9': '⁹',
            
            # Chỉ số dưới đơn giản (sử dụng subscript Unicode)
            '_0': '₀', '_1': '₁', '_2': '₂', '_3': '₃',
            '_4': '₄', '_5': '₅', '_6': '₆', '_7': '₇',
            '_8': '₈', '_9': '₉',
        }
        
        # Thực hiện chuyển đổi
        result = latex_content
        for latex_symbol, unicode_symbol in latex_to_unicode.items():
            result = result.replace(latex_symbol, unicode_symbol)
        
        # Xử lý phân số phức tạp \\frac{a}{b} -> a/b
        frac_pattern = r'\\frac\{([^}]+)\}\{([^}]+)\}'
        result = re.sub(frac_pattern, r'(\1)/(\2)', result)
        
        # Xử lý căn bậc hai \\sqrt{x} -> √x
        sqrt_pattern = r'\\sqrt\{([^}]+)\}'
        result = re.sub(sqrt_pattern, r'√(\1)', result)
        
        # Xử lý lũy thừa phức tạp {x}^{y} -> x^y
        pow_pattern = r'\{([^}]+)\}\^\{([^}]+)\}'
        result = re.sub(pow_pattern, r'\1^(\2)', result)
        
        # Xử lý chỉ số dưới phức tạp {x}_{y} -> x_y
        sub_pattern = r'\{([^}]+)\}_\{([^}]+)\}'
        result = re.sub(sub_pattern, r'\1_(\2)', result)
        
        # Loại bỏ các dấu ngoặc nhọn còn lại
        result = result.replace('{', '').replace('}', '')
        
        return result
    
    @staticmethod
    def _insert_figure_to_word(doc, tag_line, extracted_figures):
        """
        Chèn hình ảnh vào Word - xử lý cả override info và Mistral boost
        """
        try:
            # Extract figure name - xử lý cả trường hợp có override info và Mistral boost
            fig_name = None
            if 'HÌNH:' in tag_line:
                # Lấy phần sau "HÌNH:" và trước "]"
                hình_part = tag_line.split('HÌNH:')[1]
                # Loại bỏ phần override info và Mistral boost nếu có
                if '(' in hình_part:
                    fig_name = hình_part.split('(')[0].strip()
                else:
                    fig_name = hình_part.split(']')[0].strip()
            elif 'BẢNG:' in tag_line:
                # Lấy phần sau "BẢNG:" và trước "]"
                bảng_part = tag_line.split('BẢNG:')[1]
                # Loại bỏ phần override info và Mistral boost nếu có
                if '(' in bảng_part:
                    fig_name = bảng_part.split('(')[0].strip()
                else:
                    fig_name = bảng_part.split(']')[0].strip()
            
            if not fig_name or not extracted_figures:
                # Thêm placeholder text nếu không tìm thấy figure
                para = doc.add_paragraph(f"[Không tìm thấy figure: {fig_name if fig_name else 'unknown'}]")
                para.alignment = 1
                return
            
            # Tìm figure matching
            target_figure = None
            for fig in extracted_figures:
                if fig['name'] == fig_name:
                    target_figure = fig
                    break
            
            if target_figure:
                # Decode và chèn ảnh
                try:
                    img_data = base64.b64decode(target_figure['base64'])
                    img_pil = Image.open(io.BytesIO(img_data))
                    
                    # Chuyển đổi format nếu cần
                    if img_pil.mode in ('RGBA', 'LA', 'P'):
                        img_pil = img_pil.convert('RGB')
                    
                    # Tạo file tạm
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                        img_pil.save(tmp_file.name, 'PNG')
                        
                        try:
                            # Tính toán kích thước ảnh phù hợp
                            page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                            img_width = min(page_width * 0.8, Inches(6))
                        except:
                            img_width = Inches(5)
                        
                        # Chèn ảnh vào document
                        para = doc.add_paragraph()
                        para.alignment = 1  # Center alignment
                        run = para.add_run()
                        run.add_picture(tmp_file.name, width=img_width)
                        
                        # Thêm caption nếu có override info hoặc Mistral boost
                        caption_parts = []
                        if target_figure.get('override_reason'):
                            caption_parts.append(f"kept: {target_figure['override_reason']}")
                        if target_figure.get('mistral_boost'):
                            caption_parts.append(f"🧠 Mistral: {target_figure['mistral_boost']}")
                        
                        if caption_parts:
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = 1
                            caption_run = caption_para.add_run(f"({', '.join(caption_parts)})")
                            caption_run.font.size = Pt(10)
                            caption_run.font.italic = True
                        
                        # Xóa file tạm
                        os.unlink(tmp_file.name)
                    
                except Exception as img_error:
                    # Nếu lỗi xử lý ảnh, thêm placeholder
                    para = doc.add_paragraph(f"[Lỗi hiển thị {target_figure['name']}: {str(img_error)}]")
                    para.alignment = 1
            else:
                # Không tìm thấy figure matching
                para = doc.add_paragraph(f"[Không tìm thấy figure: {fig_name}]")
                para.alignment = 1
                    
        except Exception as e:
            # Lỗi parsing tag
            para = doc.add_paragraph(f"[Lỗi xử lý figure tag: {str(e)}]")
            para.alignment = 1

def display_beautiful_figures(figures, debug_img=None):
    """
    Hiển thị figures đẹp với error handling
    """
    try:
        if not figures:
            st.warning("⚠️ Không có figures nào")
            return
        
        if debug_img:
            st.image(debug_img, caption="Debug visualization", use_column_width=True)
        
        # Hiển thị figures trong grid
        cols_per_row = 3
        for i in range(0, len(figures), cols_per_row):
            cols = st.columns(cols_per_row)
            for j in range(cols_per_row):
                if i + j < len(figures):
                    fig = figures[i + j]
                    with cols[j]:
                        try:
                            img_data = base64.b64decode(fig['base64'])
                            img_pil = Image.open(io.BytesIO(img_data))
                            
                            st.image(img_pil, use_column_width=True)
                            
                            confidence_color = "🟢" if fig['confidence'] > 70 else "🟡" if fig['confidence'] > 50 else "🔴"
                            type_icon = "📊" if fig['is_table'] else "🖼️"
                            
                            override_text = ""
                            if fig.get('override_reason'):
                                override_text = f"<br><small>✅ Kept: {fig['override_reason']}</small>"
                            
                            mistral_text = ""
                            if fig.get('mistral_boost'):
                                mistral_text = f"<br><small>🧠 Mistral: {fig['mistral_boost']}</small>"
                            
                            st.markdown(f"""
                            <div style="background: #f0f0f0; padding: 0.5rem; border-radius: 5px; margin: 5px 0;">
                                <strong>{type_icon} {fig['name']}</strong><br>
                                {confidence_color} {fig['confidence']:.1f}% | {fig['method']}{override_text}{mistral_text}
                            </div>
                            """, unsafe_allow_html=True)
                        except Exception as e:
                            st.error(f"Lỗi hiển thị figure: {str(e)}")
    except Exception as e:
        st.error(f"Lỗi hiển thị figures: {str(e)}")

def validate_api_key(api_key: str) -> bool:
    if not api_key or len(api_key) < 20:
        return False
    return re.match(r'^[A-Za-z0-9_-]+, api_key) is not None

def validate_mistral_api_key(api_key: str) -> bool:
    if not api_key or len(api_key) < 20:
        return False
    # Mistral API keys usually start with specific patterns
    return True  # Simple validation for now

def format_file_size(size_bytes: int) -> str:
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024
        i += 1
    
    return f"{size_bytes:.1f} {size_names[i]}"

def clean_session_state():
    """Clean up session state to prevent memory issues"""
    keys_to_clean = [
        'pdf_latex_content', 'pdf_images', 'pdf_extracted_figures',
        'single_latex_content', 'single_extracted_figures',
        'phone_latex_content', 'phone_extracted_figures', 'phone_processed_image'
    ]
    for key in keys_to_clean:
        if key in st.session_state:
            del st.session_state[key]
    gc.collect()

def main():
    try:
        st.markdown('<h1 class="main-header">📝 PDF/LaTeX Converter - Enhanced with Mistral OCR & Phone Processing</h1>', unsafe_allow_html=True)
        
        # Hero section
        st.markdown("""
        <div style="background: linear-gradient(135deg, #FF6B35 0%, #FF8E53 100%); color: white; padding: 2rem; border-radius: 15px; margin-bottom: 2rem; text-align: center;">
            <h2 style="margin: 0;">⚖️ BALANCED TEXT FILTER + 🧠 MISTRAL OCR + 📱 PHONE PROCESSING + 📄 WORD EXPORT</h2>
            <p style="margin: 1rem 0; font-size: 1.1rem;">✅ Mistral Pixtral-12B vision analysis • ✅ Phone image processing • ✅ Word export with images • ✅ Advanced filtering</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Sidebar
        with st.sidebar:
            st.header("⚙️ Cài đặt")
            
            # Clean session button
            if st.button("🧹 Clean Memory", help="Xóa cache để giải phóng bộ nhớ"):
                clean_session_state()
                st.success("✅ Memory cleaned!")
            
            # API keys
            api_key = st.text_input("Gemini API Key", type="password")
            
            if api_key:
                if validate_api_key(api_key):
                    st.success("✅ Gemini API key hợp lệ")
                else:
                    st.error("❌ Gemini API key không hợp lệ")
            
            st.markdown("---")
            
            # Mistral OCR Service Settings
            st.markdown("### 🧠 Mistral OCR Service")
            enable_mistral_ocr = st.checkbox("Bật Mistral OCR để đếm figures", value=True)
            
            if enable_mistral_ocr:
                mistral_api_key = st.text_input(
                    "Mistral API Key", 
                    type="password",
                    help="API key cho Mistral AI service"
                )
                
                if mistral_api_key:
                    if validate_mistral_api_key(mistral_api_key):
                        st.success("✅ Mistral API key đã nhập")
                    else:
                        st.error("❌ Mistral API key quá ngắn")
                
                st.markdown("""
                <div class="mistral-badge">
                🧠 <strong>Mistral OCR Features:</strong><br>
                • Pixtral-12B vision model analysis<br>
                • Intelligent figure/table counting<br>
                • Visual complexity assessment<br>
                • Mathematical content detection<br>
                • Content type classification<br>
                • Advanced layout analysis<br>
                • Fallback to traditional method nếu lỗi
                </div>
                """, unsafe_allow_html=True)
            else:
                mistral_api_key = None
            
            st.markdown("---")
            
            # Cài đặt tách ảnh
            if CV2_AVAILABLE:
                st.markdown("### ⚖️ Balanced Text Filter")
                enable_extraction = st.checkbox("Bật tách ảnh Balanced", value=True)
                
                if enable_extraction:
                    debug_mode = st.checkbox("Debug mode", value=False)
                    
                    with st.expander("🔧 Cài đặt Advanced"):
                        confidence_threshold = st.slider("Final Confidence Threshold (%)", 50, 95, 65, 5)
                        max_figures = st.slider("Max figures per page", 5, 50, 25, 5)
                        
                        st.markdown("**Memory Management:**")
                        max_image_size = st.slider("Max image dimension", 1000, 4000, 2000, 500)
                        st.markdown(f"<small>Images larger than {max_image_size}x{max_image_size} will be resized</small>", unsafe_allow_html=True)
                        
                        st.markdown("**Word Export:**")
                        show_override_info = st.checkbox("Hiển thị override info trong Word", value=False)
                        st.markdown("<small>ℹ️ Nếu tắt, chỉ hiển thị [🖼️ HÌNH: figure-1.jpeg] thôi</small>", unsafe_allow_html=True)
            else:
                enable_extraction = False
                debug_mode = False
                st.error("❌ OpenCV không khả dụng!")
        
        if not api_key:
            st.warning("⚠️ Vui lòng nhập Gemini API Key!")
            return
        
        if not validate_api_key(api_key):
            st.error("❌ Gemini API key không hợp lệ!")
            return
        
        # Khởi tạo với error handling
        try:
            gemini_api = GeminiAPI(api_key)
            
            # Initialize Mistral OCR Service
            mistral_ocr_service = None
            if enable_mistral_ocr and mistral_api_key:
                try:
                    mistral_ocr_service = MistralOCRService(mistral_api_key)
                    st.markdown("""
                    <div class="mistral-badge">
                        🧠 Mistral OCR Service initialized with Pixtral-12B
                    </div>
                    """, unsafe_allow_html=True)
                except Exception as e:
                    st.warning(f"⚠️ Could not initialize Mistral OCR service: {str(e)}")
            elif enable_mistral_ocr:
                st.warning("⚠️ Mistral OCR enabled but missing API Key")
            
            if enable_extraction and CV2_AVAILABLE:
                image_extractor = SuperEnhancedImageExtractor(mistral_ocr_service)
                
                # Apply settings
                if 'confidence_threshold' in locals():
                    image_extractor.final_confidence_threshold = confidence_threshold
                if 'max_figures' in locals():
                    image_extractor.max_figures = max_figures
                if 'debug_mode' in locals():
                    image_extractor.debug_mode = debug_mode
                    image_extractor.content_filter.text_filter.debug_mode = debug_mode
                
                # Enable/disable OCR counting
                if mistral_ocr_service:
                    image_extractor.content_filter.enable_ocr_counting = True
                else:
                    image_extractor.content_filter.enable_ocr_counting = False
            else:
                image_extractor = None
                
        except Exception as e:
            st.error(f"❌ Lỗi khởi tạo: {str(e)}")
            return
        
        # Main content với tabs
        tab1, tab2, tab3 = st.tabs(["📄 PDF sang LaTeX", "🖼️ Ảnh sang LaTeX", "📱 Ảnh điện thoại"])
        
        with tab1:
            st.header("📄 Chuyển đổi PDF sang LaTeX")
            
            uploaded_pdf = st.file_uploader("Chọn file PDF", type=['pdf'])
            
            if uploaded_pdf:
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.subheader("📋 Preview PDF")
                    
                    # File info
                    file_size = format_file_size(uploaded_pdf.size)
                    st.info(f"📁 {uploaded_pdf.name} | 📏 {file_size}")
                    
                    # Check file size
                    if uploaded_pdf.size > 50 * 1024 * 1024:  # 50MB
                        st.warning("⚠️ File lớn (>50MB). Có thể xử lý chậm.")
                    
                    # Page limit option
                    max_pages = st.number_input("Giới hạn số trang (0 = không giới hạn)", 
                                              min_value=0, max_value=100, value=0)
                    
                    with st.spinner("🔄 Đang xử lý PDF..."):
                        try:
                            pdf_images = PDFProcessor.extract_images_and_text(
                                uploaded_pdf, 
                                max_pages if max_pages > 0 else None
                            )
                            st.success(f"✅ Đã trích xuất {len(pdf_images)} trang")
                            
                            # Preview
                            for i, (img, page_num) in enumerate(pdf_images[:2]):
                                st.markdown(f"**📄 Trang {page_num}:**")
                                st.image(img, use_column_width=True)
                            
                            if len(pdf_images) > 2:
                                st.info(f"... và {len(pdf_images) - 2} trang khác")
                        
                        except Exception as e:
                            st.error(f"❌ Lỗi xử lý PDF: {str(e)}")
                            pdf_images = []
                
                with col2:
                    st.subheader("⚡ Chuyển đổi sang LaTeX")
                    
                    if st.button("🚀 Bắt đầu chuyển đổi PDF", type="primary"):
                        if pdf_images:
                            all_latex_content = []
                            all_extracted_figures = []
                            all_debug_images = []
                            
                            # Continuous numbering across pages
                            continuous_img_idx = 0
                            continuous_table_idx = 0
                            
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            for i, (img, page_num) in enumerate(pdf_images):
                                try:
                                    status_text.text(f"Đang xử lý trang {page_num}/{len(pdf_images)}...")
                                    
                                    img_buffer = io.BytesIO()
                                    img.save(img_buffer, format='PNG')
                                    img_bytes = img_buffer.getvalue()
                                    
                                    # Check image size
                                    if len(img_bytes) > 20 * 1024 * 1024:  # 20MB
                                        st.warning(f"⚠️ Trang {page_num} quá lớn, resize...")
                                        img_resized = img.copy()
                                        img_resized.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
                                        img_buffer = io.BytesIO()
                                        img_resized.save(img_buffer, format='PNG')
                                        img_bytes = img_buffer.getvalue()
                                    
                                    # Tách ảnh với Balanced Text Filter và continuous numbering
                                    extracted_figures = []
                                    debug_img = None
                                    
                                    if enable_extraction and CV2_AVAILABLE and image_extractor:
                                        try:
                                            figures, h, w, continuous_img_idx, continuous_table_idx = image_extractor.extract_figures_and_tables(
                                                img_bytes, continuous_img_idx, continuous_table_idx
                                            )
                                            extracted_figures = figures
                                            all_extracted_figures.extend(figures)
                                            
                                            if figures:
                                                debug_img = image_extractor.create_beautiful_debug_visualization(img_bytes, figures)
                                                all_debug_images.append((debug_img, page_num, figures))
                                            
                                        except Exception as e:
                                            st.error(f"❌ Lỗi tách ảnh trang {page_num}: {str(e)}")
                                    
                                    # Prompt
                                    prompt_text = """
Chuyển đổi TOÀN BỘ nội dung trong ảnh thành văn bản với format LaTeX chính xác.

🎯 YÊU CẦU ĐỊNH DẠNG:

1. **Câu hỏi trắc nghiệm:**
```
Câu X: [nội dung câu hỏi đầy đủ]
A) [đáp án A hoàn chỉnh]
B) [đáp án B hoàn chỉnh]
C) [đáp án C hoàn chỉnh]  
D) [đáp án D hoàn chỉnh]
```

2. **Công thức toán học - LUÔN dùng ${...}$:**
- ${x^2 + y^2 = z^2}$, ${\\frac{a+b}{c-d}}$
- ${\\int_{0}^{1} x^2 dx}$, ${\\lim_{x \\to 0} \\frac{\\sin x}{x}}$
- Ví dụ: Trong hình hộp ${ABCD.A'B'C'D'}$ có tất cả các cạnh đều bằng nhau...

3. **📊 Bảng dữ liệu - Format linh hoạt:**
```
Option 1 (Multi-line):
Thời gian (phút) | [20; 25) | [25; 30) | [30; 35) | [35; 40) | [40; 45)
Số ngày | 6 | 6 | 4 | 1 | 1

Option 2 (Single-line):
Thời gian (phút) | [20; 25) | [25; 30) | [30; 35) | [35; 40) | [40; 45) Số ngày | 6 | 6 | 4 | 1 | 1
```

⚠️ TUYỆT ĐỐI dùng ${...}$ cho MỌI công thức, biến số, ký hiệu toán học!
Ví dụ: Điểm ${A}$, ${B}$, ${C}$, công thức ${x^2 + 1}$, tỉ số ${\\frac{a}{b}}$

📊 TUYỆT ĐỐI dùng | để phân cách các cột trong bảng!
Ví dụ: Tên | Tuổi | Điểm

🔹 CHÚ Ý: Chỉ dùng ký tự $ khi có cặp ${...}$, không dùng $ đơn lẻ!
""" ):
                # Đây là công thức LaTeX
                # Loại bỏ ${ và }$ để lấy nội dung bên trong
                formula_content = part[2:-2]
                
                # Chuyển đổi một số ký hiệu LaTeX cơ bản thành Unicode
                formula_content = EnhancedWordExporter._convert_latex_to_unicode(formula_content)
                
                # Thêm công thức vào paragraph với font khác biệt
                run = para.add_run(formula_content)
                run.font.name = 'Cambria Math'  # Font phù hợp cho toán học
                run.font.italic = True  # In nghiêng cho công thức
                
            elif part.strip():
                # Đây là text thường
                run = para.add_run(part)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
    
    @staticmethod
    def _convert_latex_to_unicode(latex_content):
        """
        Chuyển đổi một số ký hiệu LaTeX sang Unicode
        """
        # Dictionary chuyển đổi LaTeX sang Unicode
        latex_to_unicode = {
            # Chữ Hy Lạp
            '\\alpha': 'α', '\\beta': 'β', '\\gamma': 'γ', '\\delta': 'δ',
            '\\epsilon': 'ε', '\\theta': 'θ', '\\lambda': 'λ', '\\mu': 'μ',
            '\\pi': 'π', '\\sigma': 'σ', '\\phi': 'φ', '\\omega': 'ω',
            '\\Delta': 'Δ', '\\Theta': 'Θ', '\\Lambda': 'Λ', '\\Pi': 'Π',
            '\\Sigma': 'Σ', '\\Phi': 'Φ', '\\Omega': 'Ω',
            
            # Ký hiệu toán học
            '\\infty': '∞', '\\pm': '±', '\\mp': '∓',
            '\\times': '×', '\\div': '÷', '\\cdot': '·',
            '\\leq': '≤', '\\geq': '≥', '\\neq': '≠',
            '\\approx': '≈', '\\equiv': '≡', '\\sim': '∼',
            '\\subset': '⊂', '\\supset': '⊃', '\\in': '∈',
            '\\notin': '∉', '\\cup': '∪', '\\cap': '∩',
            '\\sum': '∑', '\\prod': '∏', '\\int': '∫',
            '\\partial': '∂', '\\nabla': '∇',
            
            # Mũi tên
            '\\rightarrow': '→', '\\leftarrow': '←',
            '\\leftrightarrow': '↔', '\\Rightarrow': '⇒',
            '\\Leftarrow': '⇐', '\\Leftrightarrow': '⇔',
            
            # Xử lý phân số đơn giản
            '\\frac{1}{2}': '½', '\\frac{1}{3}': '⅓', '\\frac{2}{3}': '⅔',
            '\\frac{1}{4}': '¼', '\\frac{3}{4}': '¾', '\\frac{1}{8}': '⅛',
            
            # Lũy thừa đơn giản (sử dụng superscript Unicode)
            '^2': '²', '^3': '³', '^1': '¹',
            '^0': '⁰', '^4': '⁴', '^5': '⁵',
            '^6': '⁶', '^7': '⁷', '^8': '⁸', '^9': '⁹',
            
            # Chỉ số dưới đơn giản (sử dụng subscript Unicode)
            '_0': '₀', '_1': '₁', '_2': '₂', '_3': '₃',
            '_4': '₄', '_5': '₅', '_6': '₆', '_7': '₇',
            '_8': '₈', '_9': '₉',
        }
        
        # Thực hiện chuyển đổi
        result = latex_content
        for latex_symbol, unicode_symbol in latex_to_unicode.items():
            result = result.replace(latex_symbol, unicode_symbol)
        
        # Xử lý phân số phức tạp \\frac{a}{b} -> a/b
        frac_pattern = r'\\frac\{([^}]+)\}\{([^}]+)\}'
        result = re.sub(frac_pattern, r'(\1)/(\2)', result)
        
        # Xử lý căn bậc hai \\sqrt{x} -> √x
        sqrt_pattern = r'\\sqrt\{([^}]+)\}'
        result = re.sub(sqrt_pattern, r'√(\1)', result)
        
        # Xử lý lũy thừa phức tạp {x}^{y} -> x^y
        pow_pattern = r'\{([^}]+)\}\^\{([^}]+)\}'
        result = re.sub(pow_pattern, r'\1^(\2)', result)
        
        # Xử lý chỉ số dưới phức tạp {x}_{y} -> x_y
        sub_pattern = r'\{([^}]+)\}_\{([^}]+)\}'
        result = re.sub(sub_pattern, r'\1_(\2)', result)
        
        # Loại bỏ các dấu ngoặc nhọn còn lại
        result = result.replace('{', '').replace('}', '')
        
        return result
    
    @staticmethod
    def _insert_figure_to_word(doc, tag_line, extracted_figures):
        """
        Chèn hình ảnh vào Word - xử lý cả override info và Mistral boost
        """
        try:
            # Extract figure name - xử lý cả trường hợp có override info và Mistral boost
            fig_name = None
            if 'HÌNH:' in tag_line:
                # Lấy phần sau "HÌNH:" và trước "]"
                hình_part = tag_line.split('HÌNH:')[1]
                # Loại bỏ phần override info và Mistral boost nếu có
                if '(' in hình_part:
                    fig_name = hình_part.split('(')[0].strip()
                else:
                    fig_name = hình_part.split(']')[0].strip()
            elif 'BẢNG:' in tag_line:
                # Lấy phần sau "BẢNG:" và trước "]"
                bảng_part = tag_line.split('BẢNG:')[1]
                # Loại bỏ phần override info và Mistral boost nếu có
                if '(' in bảng_part:
                    fig_name = bảng_part.split('(')[0].strip()
                else:
                    fig_name = bảng_part.split(']')[0].strip()
            
            if not fig_name or not extracted_figures:
                # Thêm placeholder text nếu không tìm thấy figure
                para = doc.add_paragraph(f"[Không tìm thấy figure: {fig_name if fig_name else 'unknown'}]")
                para.alignment = 1
                return
            
            # Tìm figure matching
            target_figure = None
            for fig in extracted_figures:
                if fig['name'] == fig_name:
                    target_figure = fig
                    break
            
            if target_figure:
                # Decode và chèn ảnh
                try:
                    img_data = base64.b64decode(target_figure['base64'])
                    img_pil = Image.open(io.BytesIO(img_data))
                    
                    # Chuyển đổi format nếu cần
                    if img_pil.mode in ('RGBA', 'LA', 'P'):
                        img_pil = img_pil.convert('RGB')
                    
                    # Tạo file tạm
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
                        img_pil.save(tmp_file.name, 'PNG')
                        
                        try:
                            # Tính toán kích thước ảnh phù hợp
                            page_width = doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin
                            img_width = min(page_width * 0.8, Inches(6))
                        except:
                            img_width = Inches(5)
                        
                        # Chèn ảnh vào document
                        para = doc.add_paragraph()
                        para.alignment = 1  # Center alignment
                        run = para.add_run()
                        run.add_picture(tmp_file.name, width=img_width)
                        
                        # Thêm caption nếu có override info hoặc Mistral boost
                        caption_parts = []
                        if target_figure.get('override_reason'):
                            caption_parts.append(f"kept: {target_figure['override_reason']}")
                        if target_figure.get('mistral_boost'):
                            caption_parts.append(f"🧠 Mistral: {target_figure['mistral_boost']}")
                        
                        if caption_parts:
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = 1
                            caption_run = caption_para.add_run(f"({', '.join(caption_parts)})")
                            caption_run.font.size = Pt(10)
                            caption_run.font.italic = True
                        
                        # Xóa file tạm
                        os.unlink(tmp_file.name)
                    
                except Exception as img_error:
                    # Nếu lỗi xử lý ảnh, thêm placeholder
                    para = doc.add_paragraph(f"[Lỗi hiển thị {target_figure['name']}: {str(img_error)}]")
                    para.alignment = 1
            else:
                # Không tìm thấy figure matching
                para = doc.add_paragraph(f"[Không tìm thấy figure: {fig_name}]")
                para.alignment = 1
                    
        except Exception as e:
            # Lỗi parsing tag
            para = doc.add_paragraph(f"[Lỗi xử lý figure tag: {str(e)}]")
            para.alignment = 1

def display_beautiful_figures(figures, debug_img=None):
    """
    Hiển thị figures đẹp với error handling
    """
    try:
        if not figures:
            st.warning("⚠️ Không có figures nào")
            return
        
        if debug_img:
            st.image(debug_img, caption="Debug visualization", use_column_width=True)
        
        # Hiển thị figures trong grid
        cols_per_row = 3
        for i in range(0, len(figures), cols_per_row):
            cols = st.columns(cols_per_row)
            for j in range(cols_per_row):
                if i + j < len(figures):
                    fig = figures[i + j]
                    with cols[j]:
                        try:
                            img_data = base64.b64decode(fig['base64'])
                            img_pil = Image.open(io.BytesIO(img_data))
                            
                            st.image(img_pil, use_column_width=True)
                            
                            confidence_color = "🟢" if fig['confidence'] > 70 else "🟡" if fig['confidence'] > 50 else "🔴"
                            type_icon = "📊" if fig['is_table'] else "🖼️"
                            
                            override_text = ""
                            if fig.get('override_reason'):
                                override_text = f"<br><small>✅ Kept: {fig['override_reason']}</small>"
                            
                            mistral_text = ""
                            if fig.get('mistral_boost'):
                                mistral_text = f"<br><small>🧠 Mistral: {fig['mistral_boost']}</small>"
                            
                            st.markdown(f"""
                            <div style="background: #f0f0f0; padding: 0.5rem; border-radius: 5px; margin: 5px 0;">
                                <strong>{type_icon} {fig['name']}</strong><br>
                                {confidence_color} {fig['confidence']:.1f}% | {fig['method']}{override_text}{mistral_text}
                            </div>
                            """, unsafe_allow_html=True)
                        except Exception as e:
                            st.error(f"Lỗi hiển thị figure: {str(e)}")
    except Exception as e:
        st.error(f"Lỗi hiển thị figures: {str(e)}")

def validate_api_key(api_key: str) -> bool:
    if not api_key or len(api_key) < 20:
        return False
    return re.match(r'^[A-Za-z0-9_-]+, api_key) is not None

def validate_mistral_api_key(api_key: str) -> bool:
    if not api_key or len(api_key) < 20:
        return False
    # Mistral API keys usually start with specific patterns
    return True  # Simple validation for now

def format_file_size(size_bytes: int) -> str:
    if size_bytes == 0:
        return "0 B"
    
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024
        i += 1
    
    return f"{size_bytes:.1f} {size_names[i]}"

def clean_session_state():
    """Clean up session state to prevent memory issues"""
    keys_to_clean = [
        'pdf_latex_content', 'pdf_images', 'pdf_extracted_figures',
        'single_latex_content', 'single_extracted_figures',
        'phone_latex_content', 'phone_extracted_figures', 'phone_processed_image'
    ]
    for key in keys_to_clean:
        if key in st.session_state:
            del st.session_state[key]
    gc.collect()

def main():
    try:
        st.markdown('<h1 class="main-header">📝 PDF/LaTeX Converter - Enhanced with Mistral OCR & Phone Processing</h1>', unsafe_allow_html=True)
        
        # Hero section
        st.markdown("""
        <div style="background: linear-gradient(135deg, #FF6B35 0%, #FF8E53 100%); color: white; padding: 2rem; border-radius: 15px; margin-bottom: 2rem; text-align: center;">
            <h2 style="margin: 0;">⚖️ BALANCED TEXT FILTER + 🧠 MISTRAL OCR + 📱 PHONE PROCESSING + 📄 WORD EXPORT</h2>
            <p style="margin: 1rem 0; font-size: 1.1rem;">✅ Mistral Pixtral-12B vision analysis • ✅ Phone image processing • ✅ Word export with images • ✅ Advanced filtering</p>
        </div>
        """, unsafe_allow_html=True)
        
        # Sidebar
        with st.sidebar:
            st.header("⚙️ Cài đặt")
            
            # Clean session button
            if st.button("🧹 Clean Memory", help="Xóa cache để giải phóng bộ nhớ"):
                clean_session_state()
                st.success("✅ Memory cleaned!")
            
            # API keys
            api_key = st.text_input("Gemini API Key", type="password")
            
            if api_key:
                if validate_api_key(api_key):
                    st.success("✅ Gemini API key hợp lệ")
                else:
                    st.error("❌ Gemini API key không hợp lệ")
            
            st.markdown("---")
            
            # Mistral OCR Service Settings
            st.markdown("### 🧠 Mistral OCR Service")
            enable_mistral_ocr = st.checkbox("Bật Mistral OCR để đếm figures", value=True)
            
            if enable_mistral_ocr:
                mistral_api_key = st.text_input(
                    "Mistral API Key", 
                    type="password",
                    help="API key cho Mistral AI service"
                )
                
                if mistral_api_key:
                    if validate_mistral_api_key(mistral_api_key):
                        st.success("✅ Mistral API key đã nhập")
                    else:
                        st.error("❌ Mistral API key quá ngắn")
                
                st.markdown("""
                <div class="mistral-badge">
                🧠 <strong>Mistral OCR Features:</strong><br>
                • Pixtral-12B vision model analysis<br>
                • Intelligent figure/table counting<br>
                • Visual complexity assessment<br>
                • Mathematical content detection<br>
                • Content type classification<br>
                • Advanced layout analysis<br>
                • Fallback to traditional method nếu lỗi
                </div>
                """, unsafe_allow_html=True)
            else:
                mistral_api_key = None
            
            st.markdown("---")
            
            # Cài đặt tách ảnh
            if CV2_AVAILABLE:
                st.markdown("### ⚖️ Balanced Text Filter")
                enable_extraction = st.checkbox("Bật tách ảnh Balanced", value=True)
                
                if enable_extraction:
                    debug_mode = st.checkbox("Debug mode", value=False)
                    
                    with st.expander("🔧 Cài đặt Advanced"):
                        confidence_threshold = st.slider("Final Confidence Threshold (%)", 50, 95, 65, 5)
                        max_figures = st.slider("Max figures per page", 5, 50, 25, 5)
                        
                        st.markdown("**Memory Management:**")
                        max_image_size = st.slider("Max image dimension", 1000, 4000, 2000, 500)
                        st.markdown(f"<small>Images larger than {max_image_size}x{max_image_size} will be resized</small>", unsafe_allow_html=True)
                        
                        st.markdown("**Word Export:**")
                        show_override_info = st.checkbox("Hiển thị override info trong Word", value=False)
                        st.markdown("<small>ℹ️ Nếu tắt, chỉ hiển thị [🖼️ HÌNH: figure-1.jpeg] thôi</small>", unsafe_allow_html=True)
            else:
                enable_extraction = False
                debug_mode = False
                st.error("❌ OpenCV không khả dụng!")
        
        if not api_key:
            st.warning("⚠️ Vui lòng nhập Gemini API Key!")
            return
        
        if not validate_api_key(api_key):
            st.error("❌ Gemini API key không hợp lệ!")
            return
        
        # Khởi tạo với error handling
        try:
            gemini_api = GeminiAPI(api_key)
            
            # Initialize Mistral OCR Service
            mistral_ocr_service = None
            if enable_mistral_ocr and mistral_api_key:
                try:
                    mistral_ocr_service = MistralOCRService(mistral_api_key)
                    st.markdown("""
                    <div class="mistral-badge">
                        🧠 Mistral OCR Service initialized with Pixtral-12B
                    </div>
                    """, unsafe_allow_html=True)
                except Exception as e:
                    st.warning(f"⚠️ Could not initialize Mistral OCR service: {str(e)}")
            elif enable_mistral_ocr:
                st.warning("⚠️ Mistral OCR enabled but missing API Key")
            
            if enable_extraction and CV2_AVAILABLE:
                image_extractor = SuperEnhancedImageExtractor(mistral_ocr_service)
                
                # Apply settings
                if 'confidence_threshold' in locals():
                    image_extractor.final_confidence_threshold = confidence_threshold
                if 'max_figures' in locals():
                    image_extractor.max_figures = max_figures
                if 'debug_mode' in locals():
                    image_extractor.debug_mode = debug_mode
                    image_extractor.content_filter.text_filter.debug_mode = debug_mode
                
                # Enable/disable OCR counting
                if mistral_ocr_service:
                    image_extractor.content_filter.enable_ocr_counting = True
                else:
                    image_extractor.content_filter.enable_ocr_counting = False
            else:
                image_extractor = None
                
        except Exception as e:
            st.error(f"❌ Lỗi khởi tạo: {str(e)}")
            return
        
        # Main content với tabs
        tab1, tab2, tab3 = st.tabs(["📄 PDF sang LaTeX", "🖼️ Ảnh sang LaTeX", "📱 Ảnh điện thoại"])
        
        with tab1:
            st.header("📄 Chuyển đổi PDF sang LaTeX")
            
            uploaded_pdf = st.file_uploader("Chọn file PDF", type=['pdf'])
            
            if uploaded_pdf:
                col1, col2 = st.columns([1, 1])
                
                with col1:
                    st.subheader("📋 Preview PDF")
                    
                    # File info
                    file_size = format_file_size(uploaded_pdf.size)
                    st.info(f"📁 {uploaded_pdf.name} | 📏 {file_size}")
                    
                    # Check file size
                    if uploaded_pdf.size > 50 * 1024 * 1024:  # 50MB
                        st.warning("⚠️ File lớn (>50MB). Có thể xử lý chậm.")
                    
                    # Page limit option
                    max_pages = st.number_input("Giới hạn số trang (0 = không giới hạn)", 
                                              min_value=0, max_value=100, value=0)
                    
                    with st.spinner("🔄 Đang xử lý PDF..."):
                        try:
                            pdf_images = PDFProcessor.extract_images_and_text(
                                uploaded_pdf, 
                                max_pages if max_pages > 0 else None
                            )
                            st.success(f"✅ Đã trích xuất {len(pdf_images)} trang")
                            
                            # Preview
                            for i, (img, page_num) in enumerate(pdf_images[:2]):
                                st.markdown(f"**📄 Trang {page_num}:**")
                                st.image(img, use_column_width=True)
                            
                            if len(pdf_images) > 2:
                                st.info(f"... và {len(pdf_images) - 2} trang khác")
                        
                        except Exception as e:
                            st.error(f"❌ Lỗi xử lý PDF: {str(e)}")
                            pdf_images = []
                
                with col2:
                    st.subheader("⚡ Chuyển đổi sang LaTeX")
                    
                    if st.button("🚀 Bắt đầu chuyển đổi PDF", type="primary"):
                        if pdf_images:
                            all_latex_content = []
                            all_extracted_figures = []
                            all_debug_images = []
                            
                            # Continuous numbering across pages
                            continuous_img_idx = 0
                            continuous_table_idx = 0
                            
                            progress_bar = st.progress(0)
                            status_text = st.empty()
                            
                            for i, (img, page_num) in enumerate(pdf_images):
                                try:
                                    status_text.text(f"Đang xử lý trang {page_num}/{len(pdf_images)}...")
                                    
                                    img_buffer = io.BytesIO()
                                    img.save(img_buffer, format='PNG')
                                    img_bytes = img_buffer.getvalue()
                                    
                                    # Check image size
                                    if len(img_bytes) > 20 * 1024 * 1024:  # 20MB
                                        st.warning(f"⚠️ Trang {page_num} quá lớn, resize...")
                                        img_resized = img.copy()
                                        img_resized.thumbnail((2000, 2000), Image.Resampling.LANCZOS)
                                        img_buffer = io.BytesIO()
                                        img_resized.save(img_buffer, format='PNG')
                                        img_bytes = img_buffer.getvalue()
                                    
                                    # Tách ảnh với Balanced Text Filter và continuous numbering
                                    extracted_figures = []
                                    debug_img = None
                                    
                                    if enable_extraction and CV2_AVAILABLE and image_extractor:
                                        try:
                                            figures, h, w, continuous_img_idx, continuous_table_idx = image_extractor.extract_figures_and_tables(
                                                img_bytes, continuous_img_idx, continuous_table_idx
                                            )
                                            extracted_figures = figures
                                            all_extracted_figures.extend(figures)
                                            
                                            if figures:
                                                debug_img = image_extractor.create_beautiful_debug_visualization(img_bytes, figures)
                                                all_debug_images.append((debug_img, page_num, figures))
                                            
                                        except Exception as e:
                                            st.error(f"❌ Lỗi tách ảnh trang {page_num}: {str(e)}")
                                    
                                    # Prompt
                                    prompt_text = """
Chuyển đổi TOÀN BỘ nội dung trong ảnh thành văn bản với format LaTeX chính xác.

🎯 YÊU CẦU ĐỊNH DẠNG:

1. **Câu hỏi trắc nghiệm:**
```
Câu X: [nội dung câu hỏi đầy đủ]
A) [đáp án A hoàn chỉnh]
B) [đáp án B hoàn chỉnh]
C) [đáp án C hoàn chỉnh]  
D) [đáp án D hoàn chỉnh]
```

2. **Công thức toán học - LUÔN dùng ${...}$:**
- ${x^2 + y^2 = z^2}$, ${\\frac{a+b}{c-d}}$
- ${\\int_{0}^{1} x^2 dx}$, ${\\lim_{x \\to 0} \\frac{\\sin x}{x}}$
- Ví dụ: Trong hình hộp ${ABCD.A'B'C'D'}$ có tất cả các cạnh đều bằng nhau...

3. **📊 Bảng dữ liệu - Format linh hoạt:**
```
Option 1 (Multi-line):
Thời gian (phút) | [20; 25) | [25; 30) | [30; 35) | [35; 40) | [40; 45)
Số ngày | 6 | 6 | 4 | 1 | 1

Option 2 (Single-line):
Thời gian (phút) | [20; 25) | [25; 30) | [30; 35) | [35; 40) | [40; 45) Số ngày | 6 | 6 | 4 | 1 | 1
```

⚠️ TUYỆT ĐỐI dùng ${...}$ cho MỌI công thức, biến số, ký hiệu toán học!
Ví dụ: Điểm ${A}$, ${B}$, ${C}$, công thức ${x^2 + 1}$, tỉ số ${\\frac{a}{b}}$

📊 TUYỆT ĐỐI dùng | để phân cách các cột trong bảng!
Ví dụ: Tên | Tuổi | Điểm

🔹 CHÚ Ý: Chỉ dùng ký tự $ khi có cặp ${...}$, không dùng $ đơn lẻ!
"""
